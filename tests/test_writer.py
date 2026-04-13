from docx import Document

from tools.report_scaffold_v3.docx_inspector import inspect_docx, inspect_docx_overrides
from tools.report_scaffold_v3.manifest_loader import (
    load_baseline_manifest,
    load_build_source,
    load_layout_contract,
    load_template_profile,
    load_write_contract,
)
from tools.report_scaffold_v3.ooxml_inspector import extract_effective_style_inventory
from tools.report_scaffold_v3.validation import build_validation_report
from tools.report_scaffold_v3.writer import build_word_candidate


def assert_pt_close(actual: float, expected: float, tolerance: float = 0.1):
    assert abs(actual - expected) <= tolerance


def test_build_word_candidate_writes_expected_paragraphs(tmp_path):
    baseline = load_baseline_manifest()
    source = load_build_source()
    contract = load_write_contract()
    output_docx = tmp_path / source.output_name

    result = build_word_candidate(
        template_docx=baseline.format_authority_docx.expanduser().resolve(),
        source=source,
        contract=contract,
        output_docx=output_docx,
    )

    assert result.output_docx == output_docx
    assert output_docx.exists()

    document = Document(str(output_docx))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assert paragraph_texts == [
        "Document Operations System",
        "Platform Writer Demo",
        "This candidate document is written from explicit platform rules rather than by imitating a historical case-study report.",
        "The writer clears the template body, applies declared styles, and produces a safe candidate output in the platform build directory.",
        "Appendix — Platform Placeholder",
        "This placeholder appendix block exists so candidate outputs can satisfy the current structural validator without touching protected reference documents.",
    ]


def test_build_word_candidate_writes_declared_styles_and_page_geometry(tmp_path):
    baseline = load_baseline_manifest()
    source = load_build_source()
    contract = load_write_contract()
    output_docx = tmp_path / source.output_name

    build_word_candidate(
        template_docx=baseline.format_authority_docx.expanduser().resolve(),
        source=source,
        contract=contract,
        output_docx=output_docx,
    )

    document = Document(str(output_docx))
    section = document.sections[0]
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]

    assert_pt_close(section.page_width.pt, 595.3)
    assert_pt_close(section.page_height.pt, 841.9)
    assert_pt_close(section.top_margin.pt, 70.85)
    assert_pt_close(section.bottom_margin.pt, 70.85)
    assert_pt_close(section.left_margin.pt, 70.85)
    assert_pt_close(section.right_margin.pt, 70.85)

    assert [paragraph.style.name for paragraph in paragraphs] == [
        "Heading 1",
        "Heading 2",
        "Normal",
        "Indented Body",
        "Heading 2",
        "Normal",
    ]
    assert paragraphs[0].alignment == 1
    assert paragraphs[1].alignment == 0
    assert paragraphs[2].alignment == 3
    assert round(paragraphs[3].paragraph_format.first_line_indent.pt, 1) == 24.0
    assert paragraphs[3].paragraph_format.line_spacing == 1.5


def test_built_candidate_validates_without_errors_or_warnings(tmp_path):
    baseline = load_baseline_manifest()
    source = load_build_source()
    contract = load_write_contract()
    template_profile = load_template_profile().model_dump(mode="json")
    layout_contract = load_layout_contract().model_dump(mode="json")
    output_docx = tmp_path / source.output_name

    build_word_candidate(
        template_docx=baseline.format_authority_docx.expanduser().resolve(),
        source=source,
        contract=contract,
        output_docx=output_docx,
    )

    report = build_validation_report(
        output_docx,
        inspect_docx(output_docx),
        target_role="candidate_output",
        format_authority_docx=baseline.format_authority_docx.expanduser().resolve(),
        template_profile=template_profile,
        layout_contract=layout_contract,
        style_inventory=extract_effective_style_inventory(output_docx),
        override_inventory=inspect_docx_overrides(output_docx),
    )

    errors = [finding for finding in report.findings if finding.severity == "error"]
    warnings = [finding for finding in report.findings if finding.severity == "warning"]

    assert not errors
    assert not warnings
