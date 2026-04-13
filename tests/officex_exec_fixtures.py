from __future__ import annotations

from pathlib import Path

from docx import Document


def create_officex_exec_fixture(
    tmp_path: Path,
    *,
    duplicate_business_model: bool = False,
) -> dict[str, Path]:
    candidate_path = tmp_path / "candidate.docx"
    document = Document()

    heading = document.add_paragraph("OfficeX Runtime Candidate")
    heading.style = "Heading 1"

    section_heading = document.add_paragraph("Business Model Analysis")
    section_heading.style = "Heading 2"
    document.add_paragraph("LoopMart-like placeholder text is not used here; this is a neutral OfficeX execution fixture.")

    document.add_paragraph(
        "To provide a structured assessment, this paragraph anchors the prefix-based review and patch flow."
    )

    if duplicate_business_model:
        duplicate_heading = document.add_paragraph("Business Model Analysis")
        duplicate_heading.style = "Heading 2"

    trace_heading = document.add_paragraph("Traceability Table")
    trace_heading.style = "Heading 2"
    table = document.add_table(rows=3, cols=5)
    headers = ["FR ID", "Requirement", "Status", "Evidence", "Notes"]
    for column_index, value in enumerate(headers):
        table.rows[0].cells[column_index].text = value
    fr07 = ["FR-07", "Search listings", "Met", "TC-17, SearchResultsPage", "Search journey works."]
    fr08 = ["FR-08", "View listing detail", "Met", "TC-18, ListingDetailPage", "Detail route works."]
    for column_index, value in enumerate(fr07):
        table.rows[1].cells[column_index].text = value
    for column_index, value in enumerate(fr08):
        table.rows[2].cells[column_index].text = value

    document.save(str(candidate_path))
    return {"candidate_path": candidate_path}
