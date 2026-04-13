from docx import Document

from tools.report_scaffold_v3.candidate_audit import (
    build_candidate_audit,
    render_candidate_audit_markdown,
)
from tools.report_scaffold_v3.manifest_loader import load_build_source


def test_build_candidate_audit_accepts_minimal_writer_demo(tmp_path):
    from tools.report_scaffold_v3.manifest_loader import load_baseline_manifest, load_write_contract
    from tools.report_scaffold_v3.writer import build_word_candidate

    baseline = load_baseline_manifest()
    source = load_build_source()
    contract = load_write_contract()
    output_docx = tmp_path / source.output_name

    build_word_candidate(
        template_docx=baseline.format_authority_docx.expanduser().resolve(),
        source=source,
        contract=contract,
        output_docx=output_docx,
    )

    report = build_candidate_audit(output_docx)

    assert report.expected_paragraph_count == 6
    assert report.actual_paragraph_count == 6
    assert report.expected_heading_count == 3
    assert report.actual_heading_count == 3
    assert report.expected_image_count == 0
    assert report.actual_figure_count == 0
    assert not report.findings


def test_build_candidate_audit_flags_duplicate_headings(tmp_path):
    docx_path = tmp_path / "bad_candidate.docx"
    document = Document()

    para = document.add_paragraph("Document Operations System")
    para.style = "Heading 1"
    para = document.add_paragraph("Platform Writer Demo")
    para.style = "Heading 2"
    para = document.add_paragraph("Platform Writer Demo")
    para.style = "Heading 2"
    para = document.add_paragraph("Body")
    para.style = "Normal"
    para = document.add_paragraph("Appendix — Platform Placeholder")
    para.style = "Heading 2"
    para = document.add_paragraph("Appendix body")
    para.style = "Normal"
    document.save(str(docx_path))

    report = build_candidate_audit(docx_path)

    codes = {finding.code for finding in report.findings}
    assert "candidate-duplicate-headings" in codes


def test_render_candidate_audit_markdown_lists_findings(tmp_path):
    docx_path = tmp_path / "bad_candidate.docx"
    document = Document()
    para = document.add_paragraph("Only body")
    para.style = "Normal"
    document.save(str(docx_path))

    report = build_candidate_audit(docx_path)
    markdown = render_candidate_audit_markdown(report)

    assert "Candidate Audit Report" in markdown
    assert "candidate-appendix-missing" in markdown


def test_build_candidate_audit_flags_figure_count_mismatch(tmp_path):
    docx_path = tmp_path / "bad_candidate_with_missing_figure.docx"
    document = Document()

    para = document.add_paragraph("Document Operations System")
    para.style = "Heading 1"
    para = document.add_paragraph("Platform Writer Demo")
    para.style = "Heading 2"
    para = document.add_paragraph("Body")
    para.style = "Normal"
    para = document.add_paragraph("Appendix — Platform Placeholder")
    para.style = "Heading 2"
    para = document.add_paragraph("Appendix body")
    para.style = "Normal"
    document.save(str(docx_path))

    source_yaml = tmp_path / "source.yml"
    source_yaml.write_text(
        "\n".join(
            [
                "schema_version: 1",
                'document_id: "candidate-with-figure"',
                'output_name: "candidate-with-figure.docx"',
                "blocks:",
                '  - kind: "paragraph"',
                '    role: "heading_1"',
                '    text: "Document Operations System"',
                '  - kind: "paragraph"',
                '    role: "heading_2"',
                '    text: "Platform Writer Demo"',
                '  - kind: "paragraph"',
                '    role: "body"',
                '    text: "Body"',
                '  - kind: "image"',
                '    role: "figure"',
                '    image_path: "/tmp/nonexistent.png"',
                '    caption: "Figure 1. Missing figure."',
                '  - kind: "paragraph"',
                '    role: "heading_2"',
                '    text: "Appendix — Platform Placeholder"',
                '  - kind: "paragraph"',
                '    role: "body"',
                '    text: "Appendix body"',
            ]
        ),
        encoding="utf-8",
    )

    report = build_candidate_audit(docx_path, build_source_path=source_yaml)

    codes = {finding.code for finding in report.findings}
    assert "candidate-figure-count-mismatch" in codes


def test_build_candidate_audit_flags_snippet_count_mismatch(tmp_path):
    docx_path = tmp_path / "bad_candidate_with_missing_snippet.docx"
    document = Document()

    para = document.add_paragraph("Document Operations System")
    para.style = "Heading 1"
    para = document.add_paragraph("Platform Writer Demo")
    para.style = "Heading 2"
    para = document.add_paragraph("Body")
    para.style = "Normal"
    para = document.add_paragraph("Appendix — Platform Placeholder")
    para.style = "Heading 2"
    para = document.add_paragraph("Appendix body")
    para.style = "Normal"
    document.save(str(docx_path))

    source_yaml = tmp_path / "source.yml"
    source_yaml.write_text(
        "\n".join(
            [
                "schema_version: 1",
                'document_id: "candidate-with-snippet"',
                'output_name: "candidate-with-snippet.docx"',
                "blocks:",
                '  - kind: "paragraph"',
                '    role: "heading_1"',
                '    text: "Document Operations System"',
                '  - kind: "paragraph"',
                '    role: "heading_2"',
                '    text: "Platform Writer Demo"',
                '  - kind: "paragraph"',
                '    role: "body"',
                '    text: "Body"',
                '  - kind: "paragraph"',
                '    role: "subtitle"',
                '    text: "Snippet SNIP-DEMO-01 - Demo Snippet (python)"',
                '  - kind: "paragraph"',
                '    role: "body"',
                '    text: "print(\\"demo\\")"',
                '  - kind: "paragraph"',
                '    role: "heading_2"',
                '    text: "Appendix — Platform Placeholder"',
                '  - kind: "paragraph"',
                '    role: "body"',
                '    text: "Appendix body"',
            ]
        ),
        encoding="utf-8",
    )

    report = build_candidate_audit(docx_path, build_source_path=source_yaml)

    codes = {finding.code for finding in report.findings}
    assert report.expected_snippet_count == 1
    assert report.actual_snippet_count == 0
    assert "candidate-snippet-count-mismatch" in codes
