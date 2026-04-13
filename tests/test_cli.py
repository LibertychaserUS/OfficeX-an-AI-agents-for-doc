import json
from pathlib import Path

import yaml
from tools.report_scaffold_v3.manifest_loader import load_baseline_manifest
from typer.testing import CliRunner

from tools.report_scaffold_v3.cli import app, resolve_target_context
from tools.report_scaffold_v3.paths import BUILD_SOURCES_DIR, BUILDS_DIR


def test_check_package_command_passes():
    runner = CliRunner()

    result = runner.invoke(app, ["check-package"])

    assert result.exit_code == 0
    assert "package integrity check passed" in result.stdout.lower()


def test_build_word_rejects_protected_output_path():
    runner = CliRunner()
    baseline = load_baseline_manifest()

    result = runner.invoke(app, ["build-word", "--output-docx", str(baseline.target_docx)])

    assert result.exit_code != 0
    normalized_output = result.output.lower()
    assert "refusing to write candidate output over protected file" in normalized_output


def test_check_fonts_command_scans_docx(tmp_path):
    from docx import Document

    runner = CliRunner()
    docx_path = tmp_path / "font_cli_sample.docx"
    output_dir = tmp_path / "font_audit"

    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Unexpected font text")
    run.font.name = "Segoe UI Emoji"
    document.save(str(docx_path))

    result = runner.invoke(
        app,
        ["check-fonts", "--docx", str(docx_path), "--output-dir", str(output_dir)],
    )

    assert result.exit_code == 0
    stdout = result.stdout.lower()
    assert "explicit other-font" in stdout
    assert "runs: 1" in stdout
    assert (output_dir / "font_audit.json").exists()


def test_check_outline_command_scans_docx(tmp_path):
    from docx import Document

    runner = CliRunner()
    docx_path = tmp_path / "outline_cli_sample.docx"
    output_dir = tmp_path / "outline_audit"

    document = Document()
    heading = document.add_paragraph("Appendix A")
    heading.style = "Heading 2"
    document.save(str(docx_path))

    result = runner.invoke(
        app,
        ["check-outline", "--docx", str(docx_path), "--output-dir", str(output_dir)],
    )

    assert result.exit_code == 0
    stdout = result.stdout.lower()
    assert "headings: 1" in stdout
    assert "appendix headings: 1" in stdout
    assert (output_dir / "outline_audit.json").exists()


def test_check_candidate_command_audits_docx(tmp_path):
    from docx import Document

    runner = CliRunner()
    docx_path = tmp_path / "candidate_cli_sample.docx"
    output_dir = tmp_path / "candidate_audit"

    document = Document()
    p = document.add_paragraph("Document Operations System")
    p.style = "Heading 1"
    p = document.add_paragraph("Platform Writer Demo")
    p.style = "Heading 2"
    p = document.add_paragraph("Body")
    p.style = "Normal"
    p = document.add_paragraph("Appendix — Platform Placeholder")
    p.style = "Heading 2"
    p = document.add_paragraph("Appendix body")
    p.style = "Normal"
    document.save(str(docx_path))

    result = runner.invoke(
        app,
        ["check-candidate", "--docx", str(docx_path), "--output-dir", str(output_dir)],
    )

    assert result.exit_code == 0
    assert "error(s)" in result.stdout.lower()
    assert (output_dir / "candidate_audit.json").exists()


def test_assemble_sections_command_writes_build_source(tmp_path):
    runner = CliRunner()
    output_source = tmp_path / "assembled.yml"
    shared_summary = BUILD_SOURCES_DIR / "assembled_sections_summary.json"
    previous_shared_summary = (
        json.loads(shared_summary.read_text(encoding="utf-8"))
        if shared_summary.exists()
        else None
    )

    result = runner.invoke(
        app,
        ["assemble-sections", "--output-source", str(output_source)],
    )

    assert result.exit_code == 0
    assert "assembled section-managed build source" in result.stdout.lower()
    assert output_source.exists()
    assert output_source.with_suffix(".summary.json").exists()
    assert output_source.with_suffix(".summary.md").exists()
    if previous_shared_summary is not None:
        assert json.loads(shared_summary.read_text(encoding="utf-8")) == previous_shared_summary


def test_build_word_command_creates_candidate_docx(tmp_path: Path):
    runner = CliRunner()
    output_docx = tmp_path / "cli_build.docx"
    shared_summary = BUILDS_DIR / "build_summary.json"
    previous_shared_summary = (
        json.loads(shared_summary.read_text(encoding="utf-8"))
        if shared_summary.exists()
        else None
    )

    result = runner.invoke(app, ["build-word", "--output-docx", str(output_docx)])

    assert result.exit_code == 0
    assert output_docx.exists()
    assert output_docx.with_suffix(".build_summary.json").exists()
    assert output_docx.with_suffix(".build_summary.md").exists()
    if previous_shared_summary is not None:
        assert json.loads(shared_summary.read_text(encoding="utf-8")) == previous_shared_summary


def test_resolve_target_context_infers_candidate_output_for_builds_dir():
    baseline = load_baseline_manifest()
    candidate_docx = BUILDS_DIR / "inferred_candidate.docx"

    resolved_docx, target_role, _ = resolve_target_context(
        candidate_docx,
        baseline_manifest_path=baseline.scaffold_root / "manifests" / "baseline.yml",
    )

    assert resolved_docx == candidate_docx.expanduser().resolve()
    assert target_role == "candidate_output"


def test_run_section_pipeline_command_creates_sequential_outputs(tmp_path: Path):
    runner = CliRunner()
    output_dir = tmp_path / "section_pipeline"

    result = runner.invoke(
        app,
        ["run-section-pipeline", "--output-dir", str(output_dir)],
    )

    assert result.exit_code == 0
    assert "ran section pipeline sequentially" in result.stdout.lower()
    assert (output_dir / "assembled_sections_build.yml").exists()
    assert (output_dir / "candidate_audit.json").exists()
    assert (output_dir / "validation.json").exists()
    assert (output_dir / "section_pipeline_summary.json").exists()


def test_check_snippets_command_audits_manifest(tmp_path: Path):
    runner = CliRunner()
    output_dir = tmp_path / "snippet_audit"
    source_path = tmp_path / "demo.py"
    source_path.write_text("alpha\nbeta\ngamma\n", encoding="utf-8")
    manifest_path = tmp_path / "snippets.yml"
    manifest_path.write_text(
        yaml.safe_dump(
            {
                "schema_version": 1,
                "managed_snippets": [
                    {
                        "snippet_id": "SNIP-DEMO-01",
                        "title": "Demo Snippet",
                        "language": "python",
                        "source_path": str(source_path),
                        "target_section_id": "assembly-demo",
                        "extract_mode": "line_range",
                        "start_line": 2,
                        "end_line": 3,
                    }
                ],
            },
            sort_keys=False,
        ),
        encoding="utf-8",
    )

    result = runner.invoke(
        app,
        [
            "check-snippets",
            "--snippets-manifest",
            str(manifest_path),
            "--output-dir",
            str(output_dir),
        ],
    )

    assert result.exit_code == 0
    assert "audited snippets" in result.stdout.lower()
    assert (output_dir / "snippet_audit.json").exists()
    assert (output_dir / "snippet_audit_summary.md").exists()


def test_publish_run_command_writes_current_metadata(tmp_path: Path):
    runner = CliRunner()
    run_dir = tmp_path / "section_pipeline"
    output_dir = tmp_path / "published"
    output_docx = run_dir / "candidate.docx"
    build_source_path = run_dir / "assembled.yml"
    candidate_audit_path = run_dir / "candidate_audit.md"
    validation_report_path = run_dir / "validation_report.md"
    run_dir.mkdir()
    output_docx.write_text("docx-placeholder", encoding="utf-8")
    build_source_path.write_text("schema_version: 1\n", encoding="utf-8")
    candidate_audit_path.write_text("# Candidate Audit\n", encoding="utf-8")
    validation_report_path.write_text("# Validation Report\n", encoding="utf-8")
    (run_dir / "candidate_audit.json").write_text('{"findings": []}', encoding="utf-8")
    (run_dir / "validation.json").write_text('{"findings": []}', encoding="utf-8")
    (run_dir / "section_pipeline_summary.json").write_text(
        json.dumps(
            {
                "build_source_path": str(build_source_path),
                "output_docx": str(output_docx),
                "candidate_audit_path": str(candidate_audit_path),
                "validation_report_path": str(validation_report_path),
                "candidate_error_count": 0,
                "candidate_warning_count": 0,
                "validation_error_count": 0,
                "validation_warning_count": 0,
            }
        ),
        encoding="utf-8",
    )

    result = runner.invoke(
        app,
        ["publish-run", "--run-dir", str(run_dir), "--output-dir", str(output_dir)],
    )

    assert result.exit_code == 0
    assert "published run" in result.stdout.lower()
    assert (output_dir / "current.json").exists()
    assert (output_dir / "current_summary.md").exists()


def test_index_trace_command_writes_catalog(tmp_path: Path):
    runner = CliRunner()
    trace_dir = tmp_path / "trace"
    trace_dir.mkdir()
    (trace_dir / "CHECKPOINT_01.md").write_text(
        "# Checkpoint 01 - Start\n\ndate: 2026-03-31\n\nInitial.\n",
        encoding="utf-8",
    )
    (trace_dir / "CHECKPOINT_03.md").write_text(
        "# Checkpoint 03 - Resume\n\ndate: 2026-04-01\n\nBackfill note.\n",
        encoding="utf-8",
    )

    result = runner.invoke(app, ["index-trace", "--trace-dir", str(trace_dir)])

    assert result.exit_code == 0
    assert "indexed trace" in result.stdout.lower()
    assert (trace_dir / "checkpoint_catalog.json").exists()
    assert (trace_dir / "checkpoint_catalog.md").exists()
    assert (trace_dir / "checkpoint_catalog_summary.md").exists()

