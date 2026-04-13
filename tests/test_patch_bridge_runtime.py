from __future__ import annotations

import json
from pathlib import Path

from docx import Document
import pytest
from typer.testing import CliRunner

from tools.report_scaffold_v3.cli import app
from tools.report_scaffold_v3.patch_bridge_runtime import (
    apply_officex_patch_bundle,
    build_patch_spec_from_officex_bundle,
)
from tools.report_scaffold_v3.review_runtime import build_review_ledger, extract_anchors_from_review_ledger

from .officex_exec_fixtures import create_officex_exec_fixture


def _write_manual_review_input(path: Path, payload: dict) -> Path:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _heading_review_input(*, review_id: str = "patch-bridge-heading") -> dict:
    return {
        "review_id": review_id,
        "target_document_id": "candidate",
        "generated_by": "test-suite",
        "findings": [
            {
                "issue_id": "OX-PATCH-001",
                "title": "Business model heading anchor",
                "issue_kind": "content_update",
                "severity": "medium",
                "status": "open",
                "review_comment": "Use the business model heading as the patch anchor.",
                "anchor_rule": {
                    "anchor_rule_type": "paragraph_text",
                    "anchor_role": "business_model_heading",
                    "block_kind": "heading",
                    "match_mode": "exact",
                    "needle": "Business Model Analysis",
                },
            }
        ],
    }


def _build_snapshot(tmp_path: Path, *, duplicate_business_model: bool = False):
    fixture = create_officex_exec_fixture(tmp_path, duplicate_business_model=duplicate_business_model)
    review_input_path = _write_manual_review_input(
        tmp_path / "patch_bridge_review.json",
        _heading_review_input(
            review_id="patch-bridge-non-unique" if duplicate_business_model else "patch-bridge-unique"
        ),
    )
    _ledger, ledger_path = build_review_ledger(review_input_path=review_input_path)
    report = extract_anchors_from_review_ledger(
        candidate_path=fixture["candidate_path"],
        review_ledger_path=ledger_path,
    )
    snapshot = json.loads(report.anchor_snapshot_path.read_text(encoding="utf-8"))
    return fixture, snapshot, report.anchor_snapshot_path, report


def _write_patch_bundle(
    path: Path,
    *,
    run_id: str,
    target_document_id: str,
    target_anchor_id: str,
    operation_kind: str,
    allowed_scope: str = "single_paragraph",
    proposed_change: dict | None = None,
    risk_level: str = "low",
    requires_user_confirmation: bool = False,
    executor_kind: str = "ooxml_text_executor",
    approval_mode: str = "ask_every_conflict",
) -> Path:
    payload = {
        "patch_bundle_id": f"{run_id}-bundle",
        "run_id": run_id,
        "target_document_id": target_document_id,
        "generated_by": "test-suite",
        "patch_intent": "content_revision",
        "operations": [
            {
                "operation_id": f"{run_id}-op-01",
                "operation_kind": operation_kind,
                "target_anchor_id": target_anchor_id,
                "allowed_scope": allowed_scope,
                "proposed_change": proposed_change or {},
                "risk_level": risk_level,
                "requires_user_confirmation": requires_user_confirmation,
                "executor_kind": executor_kind,
                "expected_effects": ["test-effect"],
            }
        ],
        "approval_mode": approval_mode,
        "approval_requirements": [],
        "verification_requirements": ["structural"],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def test_build_patch_spec_from_officex_bundle_supports_replace_text(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "replace_text_bundle.json",
        run_id="replace-text",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated Business Model Analysis"},
        executor_kind="ooxml_text_executor",
    )

    _bundle, _anchor_snapshot, spec = build_patch_spec_from_officex_bundle(
        patch_bundle_path=bundle_path,
        candidate_path=fixture["candidate_path"],
        anchor_snapshot_path=snapshot_path,
    )

    assert spec.spec_id == "replace-text-bundle"
    assert spec.operations[0].action == "replace_paragraph_text"
    assert spec.operations[0].payload["text"] == "Updated Business Model Analysis"


def test_build_patch_spec_from_officex_bundle_supports_insert_paragraph(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "insert_paragraph_bundle.json",
        run_id="insert-paragraph",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="insert_paragraph",
        proposed_change={"text": "Inserted bridge paragraph.", "style": "Normal"},
        executor_kind="ooxml_text_executor",
    )

    _bundle, _anchor_snapshot, spec = build_patch_spec_from_officex_bundle(
        patch_bundle_path=bundle_path,
        candidate_path=fixture["candidate_path"],
        anchor_snapshot_path=snapshot_path,
    )

    assert spec.operations[0].action == "insert_paragraph_after"
    assert spec.operations[0].payload["style"] == "Normal"


def test_build_patch_spec_from_officex_bundle_supports_restyle_paragraph(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "restyle_paragraph_bundle.json",
        run_id="restyle-paragraph",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="restyle_paragraph",
        proposed_change={"style": "Heading 2"},
        executor_kind="ooxml_style_executor",
    )

    _bundle, _anchor_snapshot, spec = build_patch_spec_from_officex_bundle(
        patch_bundle_path=bundle_path,
        candidate_path=fixture["candidate_path"],
        anchor_snapshot_path=snapshot_path,
    )

    assert spec.operations[0].action == "update_paragraph_style"
    assert spec.operations[0].payload["style"] == "Heading 2"


def test_build_patch_spec_from_officex_bundle_rejects_unsupported_operation_kind(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "unsupported_operation_bundle.json",
        run_id="unsupported-operation",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="insert_figure",
        proposed_change={"figure_id": "fig-01"},
        executor_kind="asset_pipeline_executor",
    )

    with pytest.raises(ValueError, match="Unsupported OfficeX operation kind"):
        build_patch_spec_from_officex_bundle(
            patch_bundle_path=bundle_path,
            candidate_path=fixture["candidate_path"],
            anchor_snapshot_path=snapshot_path,
        )


def test_build_patch_spec_from_officex_bundle_rejects_unsupported_allowed_scope(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "unsupported_scope_bundle.json",
        run_id="unsupported-scope",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        allowed_scope="single_section",
        proposed_change={"text": "Updated text"},
        executor_kind="ooxml_text_executor",
    )

    with pytest.raises(ValueError, match="Unsupported OfficeX allowed scope"):
        build_patch_spec_from_officex_bundle(
            patch_bundle_path=bundle_path,
            candidate_path=fixture["candidate_path"],
            anchor_snapshot_path=snapshot_path,
        )


def test_build_patch_spec_from_officex_bundle_rejects_missing_anchor(tmp_path: Path):
    fixture, _snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "missing_anchor_bundle.json",
        run_id="missing-anchor",
        target_document_id="candidate",
        target_anchor_id="missing-anchor-id",
        operation_kind="replace_text",
        proposed_change={"text": "Updated text"},
        executor_kind="ooxml_text_executor",
    )

    with pytest.raises(ValueError, match="Missing anchor"):
        build_patch_spec_from_officex_bundle(
            patch_bundle_path=bundle_path,
            candidate_path=fixture["candidate_path"],
            anchor_snapshot_path=snapshot_path,
        )


def test_build_patch_spec_from_officex_bundle_rejects_non_unique_anchor(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path, duplicate_business_model=True)
    assert report.finding_count == 1
    bundle_path = _write_patch_bundle(
        tmp_path / "non_unique_anchor_bundle.json",
        run_id="non-unique-anchor",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated text"},
        executor_kind="ooxml_text_executor",
    )

    with pytest.raises(ValueError, match="non-unique"):
        build_patch_spec_from_officex_bundle(
            patch_bundle_path=bundle_path,
            candidate_path=fixture["candidate_path"],
            anchor_snapshot_path=snapshot_path,
        )


def test_build_patch_spec_from_officex_bundle_rejects_candidate_hash_mismatch(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "candidate_hash_bundle.json",
        run_id="candidate-hash",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated text"},
        executor_kind="ooxml_text_executor",
    )
    fixture["candidate_path"].write_bytes(fixture["candidate_path"].read_bytes() + b"drift")

    with pytest.raises(ValueError, match="Candidate hash mismatch"):
        build_patch_spec_from_officex_bundle(
            patch_bundle_path=bundle_path,
            candidate_path=fixture["candidate_path"],
            anchor_snapshot_path=snapshot_path,
        )


def test_apply_officex_patch_bundle_dry_run_succeeds(tmp_path: Path):
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    original_texts = [paragraph.text for paragraph in Document(str(fixture["candidate_path"])).paragraphs]
    bundle_path = _write_patch_bundle(
        tmp_path / "dry_run_bundle.json",
        run_id="dry-run",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated Business Model Analysis"},
        executor_kind="ooxml_text_executor",
    )

    bridge_report = apply_officex_patch_bundle(
        patch_bundle_path=bundle_path,
        candidate_path=fixture["candidate_path"],
        anchor_snapshot_path=snapshot_path,
        dry_run=True,
    )

    assert bridge_report.status == "validated"
    assert bridge_report.dry_run is True
    assert bridge_report.patch_spec_path.exists()
    assert bridge_report.execution_report_path.exists()
    texts_after = [paragraph.text for paragraph in Document(str(fixture["candidate_path"])).paragraphs]
    assert texts_after == original_texts


def test_officex_task_apply_patch_bundle_apply_success(tmp_path: Path):
    runner = CliRunner()
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "apply_bundle.json",
        run_id="apply-success",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="insert_paragraph",
        proposed_change={"text": "Inserted bridge paragraph.", "style": "Normal"},
        executor_kind="ooxml_text_executor",
    )

    result = runner.invoke(
        app,
        [
            "officex",
            "task",
            "apply-patch-bundle",
            "--patch-bundle",
            str(bundle_path),
            "--candidate-docx",
            str(fixture["candidate_path"]),
            "--anchor-snapshot",
            str(snapshot_path),
            "--apply",
            "--as-json",
        ],
    )

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["status"] == "applied"
    assert payload["dry_run"] is False
    texts_after = [paragraph.text for paragraph in Document(str(fixture["candidate_path"])).paragraphs]
    assert "Inserted bridge paragraph." in texts_after


def test_officex_task_apply_patch_bundle_dry_run_as_json(tmp_path: Path):
    runner = CliRunner()
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "dry_run_cli_bundle.json",
        run_id="dry-run-cli",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated Business Model Analysis"},
        executor_kind="ooxml_text_executor",
    )

    result = runner.invoke(
        app,
        [
            "officex",
            "task",
            "apply-patch-bundle",
            "--patch-bundle",
            str(bundle_path),
            "--candidate-docx",
            str(fixture["candidate_path"]),
            "--anchor-snapshot",
            str(snapshot_path),
            "--dry-run",
            "--as-json",
        ],
    )

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["status"] == "validated"
    assert payload["operation_count"] == 1


def test_officex_task_apply_patch_bundle_rejects_invalid_patch_bundle_shape(tmp_path: Path):
    runner = CliRunner()
    fixture, _snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    invalid_bundle = tmp_path / "invalid_patch_bundle.json"
    invalid_bundle.write_text("{}", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "officex",
            "task",
            "apply-patch-bundle",
            "--patch-bundle",
            str(invalid_bundle),
            "--candidate-docx",
            str(fixture["candidate_path"]),
            "--anchor-snapshot",
            str(snapshot_path),
        ],
    )

    assert result.exit_code == 1
    assert "Invalid OfficeXPatchBundle schema" in result.stdout


def test_officex_task_apply_patch_bundle_rejects_unsupported_operation(tmp_path: Path):
    runner = CliRunner()
    fixture, snapshot, snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "unsupported_cli_bundle.json",
        run_id="unsupported-cli",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="insert_figure",
        proposed_change={"figure_id": "fig-01"},
        executor_kind="asset_pipeline_executor",
    )

    result = runner.invoke(
        app,
        [
            "officex",
            "task",
            "apply-patch-bundle",
            "--patch-bundle",
            str(bundle_path),
            "--candidate-docx",
            str(fixture["candidate_path"]),
            "--anchor-snapshot",
            str(snapshot_path),
        ],
    )

    assert result.exit_code == 1
    assert "Unsupported OfficeX operation kind" in result.stdout


def test_officex_task_apply_patch_bundle_rejects_missing_anchor_snapshot(tmp_path: Path):
    runner = CliRunner()
    fixture, snapshot, _snapshot_path, report = _build_snapshot(tmp_path)
    assert report.finding_count == 0
    bundle_path = _write_patch_bundle(
        tmp_path / "missing_snapshot_bundle.json",
        run_id="missing-snapshot",
        target_document_id="candidate",
        target_anchor_id=snapshot["anchors"][0]["anchor_id"],
        operation_kind="replace_text",
        proposed_change={"text": "Updated text"},
        executor_kind="ooxml_text_executor",
    )

    result = runner.invoke(
        app,
        [
            "officex",
            "task",
            "apply-patch-bundle",
            "--patch-bundle",
            str(bundle_path),
            "--candidate-docx",
            str(fixture["candidate_path"]),
            "--anchor-snapshot",
            str(tmp_path / "missing_snapshot.json"),
        ],
    )

    assert result.exit_code == 1
    assert "File not found" in result.stdout
