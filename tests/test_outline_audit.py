from docx import Document

from tools.report_scaffold_v3.outline_audit import render_outline_audit_markdown, scan_docx_outline


def test_scan_docx_outline_detects_headings_appendix_and_duplicates(tmp_path):
    docx_path = tmp_path / "outline_sample.docx"
    document = Document()

    title = document.add_paragraph("Stage 1 Planning")
    title.style = "Heading 1"

    section = document.add_paragraph("1.1 Introduction")
    section.style = "Heading 2"

    appendix = document.add_paragraph("Appendix A")
    appendix.style = "Heading 2"

    duplicate = document.add_paragraph("1.1 Introduction")
    duplicate.style = "Heading 2"

    document.save(str(docx_path))

    report = scan_docx_outline(docx_path)

    assert report.heading_count == 4
    assert report.appendix_heading_count == 1
    assert report.heading_level_counts == {"1": 1, "2": 3}
    assert report.duplicate_heading_texts == ["1.1 Introduction"]


def test_render_outline_audit_markdown_lists_heading_inventory(tmp_path):
    docx_path = tmp_path / "outline_sample.docx"
    document = Document()
    paragraph = document.add_paragraph("Appendix A")
    paragraph.style = "Heading 2"
    document.save(str(docx_path))

    report = scan_docx_outline(docx_path)
    markdown = render_outline_audit_markdown(report)

    assert "Outline Audit Report" in markdown
    assert "Appendix A" in markdown
    assert "Heading 2" in markdown
