from docx import Document

from tools.report_scaffold_v3.font_audit import render_font_audit_markdown, scan_docx_fonts


def test_scan_docx_fonts_detects_inherited_expected_and_other_fonts(tmp_path):
    docx_path = tmp_path / "font_audit_sample.docx"
    document = Document()

    inherited_paragraph = document.add_paragraph()
    inherited_paragraph.add_run("Inherited body text")

    expected_paragraph = document.add_paragraph()
    expected_run = expected_paragraph.add_run("Explicit expected font")
    expected_run.font.name = "Times New Roman"

    violation_paragraph = document.add_paragraph()
    violation_run = violation_paragraph.add_run("Explicit unexpected font")
    violation_run.font.name = "Segoe UI Symbol"

    document.save(str(docx_path))

    report = scan_docx_fonts(docx_path)

    assert report.total_runs_scanned == 3
    assert report.inherited_font_runs == 1
    assert report.explicit_expected_font_runs == 1
    assert report.explicit_other_font_runs == 1
    assert len(report.violations) == 1
    assert report.violations[0].font_name == "Segoe UI Symbol"
    assert report.violations[0].occurrences == 1


def test_render_font_audit_markdown_lists_violation_examples(tmp_path):
    docx_path = tmp_path / "font_audit_sample.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Unexpected font text")
    run.font.name = "Segoe UI Emoji"
    document.save(str(docx_path))

    report = scan_docx_fonts(docx_path)
    markdown = render_font_audit_markdown(report)

    assert "Font Audit Report" in markdown
    assert "Segoe UI Emoji" in markdown
    assert "Paragraph 0" in markdown
