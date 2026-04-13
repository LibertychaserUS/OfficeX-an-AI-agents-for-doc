from __future__ import annotations

from pathlib import Path

from docx import Document

from .font_audit import scan_docx_fonts
from .docx_inspector import inspect_docx
from .manifest_loader import load_build_source, load_write_contract
from .models import CandidateAuditFinding, CandidateAuditReport
from .outline_audit import scan_docx_outline


def _expected_paragraph_count(source_manifest) -> int:
    total = 0
    for block in source_manifest.blocks:
        if block.get("kind") == "paragraph":
            total += 1
        elif block.get("kind") == "image":
            total += 2
    return total


def _expected_heading_count(source_manifest, write_contract_manifest) -> int:
    total = 0
    for block in source_manifest.blocks:
        if block.get("kind") != "paragraph":
            continue
        role_name = block.get("role")
        role = write_contract_manifest.paragraph_roles.get(role_name)
        if role is None:
            continue
        if role.style.lower().startswith("heading ") or role.style.lower() in {"title", "subtitle"}:
            total += 1
    return total


def _expected_image_count(source_manifest) -> int:
    return sum(1 for block in source_manifest.blocks if block.get("kind") == "image")


def _expected_snippet_count(source_manifest) -> int:
    return sum(
        1
        for block in source_manifest.blocks
        if block.get("kind") == "paragraph" and str(block.get("text", "")).startswith("Snippet ")
    )


def _allowed_styles(write_contract_manifest) -> list[str]:
    return sorted({role.style for role in write_contract_manifest.paragraph_roles.values()})


def build_candidate_audit(
    docx_path: Path,
    *,
    build_source_path: Path | None = None,
    write_contract_path: Path | None = None,
) -> CandidateAuditReport:
    source_manifest = load_build_source(build_source_path) if build_source_path else load_build_source()
    write_contract_manifest = (
        load_write_contract(write_contract_path) if write_contract_path else load_write_contract()
    )

    expected_paragraphs = _expected_paragraph_count(source_manifest)
    expected_headings = _expected_heading_count(source_manifest, write_contract_manifest)
    expected_images = _expected_image_count(source_manifest)
    expected_snippets = _expected_snippet_count(source_manifest)
    allowed_styles = _allowed_styles(write_contract_manifest)

    outline_report = scan_docx_outline(docx_path)
    font_report = scan_docx_fonts(docx_path)
    inventory = inspect_docx(docx_path)
    document = Document(str(docx_path))

    findings: list[CandidateAuditFinding] = []

    actual_paragraph_count = int(inventory["summary"]["paragraph_count"])
    actual_figure_count = int(inventory["summary"]["figure_count"])
    actual_snippet_count = sum(
        1 for paragraph in document.paragraphs if paragraph.text.strip().startswith("Snippet ")
    )

    if actual_paragraph_count != expected_paragraphs:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-paragraph-count-mismatch",
                message=(
                    f"Candidate paragraph count {actual_paragraph_count} does not match "
                    f"build-source expectation {expected_paragraphs}."
                ),
            )
        )

    if outline_report.heading_count != expected_headings:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-heading-count-mismatch",
                message=(
                    f"Candidate heading count {outline_report.heading_count} does not match "
                    f"build-source expectation {expected_headings}."
                ),
            )
        )

    if actual_figure_count != expected_images:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-figure-count-mismatch",
                message=(
                    f"Candidate figure count {actual_figure_count} does not match "
                    f"build-source expectation {expected_images}."
                ),
            )
        )

    if actual_snippet_count != expected_snippets:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-snippet-count-mismatch",
                message=(
                    f"Candidate snippet count {actual_snippet_count} does not match "
                    f"build-source expectation {expected_snippets}."
                ),
            )
        )

    used_styles = sorted(
        {
            paragraph.style.name
            for paragraph in document.paragraphs
            if paragraph.text.strip() and paragraph.style is not None
        }
    )
    unexpected_styles = sorted(set(used_styles) - set(allowed_styles))
    if unexpected_styles:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-unexpected-styles",
                message=(
                    "Candidate output uses paragraph styles outside the write contract: "
                    + ", ".join(unexpected_styles)
                ),
            )
        )

    if outline_report.appendix_heading_count == 0:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-appendix-missing",
                message="Candidate output does not contain an appendix heading.",
            )
        )

    if outline_report.duplicate_heading_texts:
        findings.append(
            CandidateAuditFinding(
                severity="error",
                code="candidate-duplicate-headings",
                message=(
                    "Candidate output contains duplicate heading texts: "
                    + ", ".join(outline_report.duplicate_heading_texts)
                ),
            )
        )

    if font_report.explicit_other_font_runs > 0:
        findings.append(
            CandidateAuditFinding(
                severity="warning",
                code="candidate-explicit-other-font-runs",
                message=(
                    f"Candidate output contains {font_report.explicit_other_font_runs} explicit non-expected-font runs."
                ),
            )
        )

    return CandidateAuditReport(
        source_docx=docx_path,
        expected_paragraph_count=expected_paragraphs,
        actual_paragraph_count=actual_paragraph_count,
        expected_heading_count=expected_headings,
        actual_heading_count=outline_report.heading_count,
        expected_image_count=expected_images,
        actual_figure_count=actual_figure_count,
        expected_snippet_count=expected_snippets,
        actual_snippet_count=actual_snippet_count,
        allowed_styles=allowed_styles,
        findings=findings,
    )


def render_candidate_audit_markdown(report: CandidateAuditReport) -> str:
    lines = [
        "# Candidate Audit Report",
        "",
        f"- Source docx: `{report.source_docx}`",
        f"- Expected paragraph count: {report.expected_paragraph_count}",
        f"- Actual paragraph count: {report.actual_paragraph_count}",
        f"- Expected heading count: {report.expected_heading_count}",
        f"- Actual heading count: {report.actual_heading_count}",
        f"- Expected image count: {report.expected_image_count}",
        f"- Actual figure count: {report.actual_figure_count}",
        f"- Expected snippet count: {report.expected_snippet_count}",
        f"- Actual snippet count: {report.actual_snippet_count}",
        f"- Allowed styles: {', '.join(report.allowed_styles)}",
        "",
        "## Findings",
        "",
    ]
    if not report.findings:
        lines.append("- No candidate-specific findings.")
        return "\n".join(lines)

    for finding in report.findings:
        lines.append(f"- [{finding.severity.upper()}] `{finding.code}`: {finding.message}")
    return "\n".join(lines)
