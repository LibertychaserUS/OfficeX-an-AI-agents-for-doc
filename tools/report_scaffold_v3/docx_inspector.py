from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from hashlib import sha1
from pathlib import Path
import re

from docx import Document

EMU_PER_PT = 12700.0
FIGURE_PATTERN = re.compile(
    r"^(Appendix Figure|Figure|Fig\.?|Table)\s+([A-Za-z0-9.\-]+)\s*[:.\-]?\s*(.*)$"
)
FILE_REFERENCE_PATTERN = re.compile(
    r"([A-Za-z0-9_./\\-]+\.(?:java|js|jsx|ts|tsx|sql|xml|json|yml|yaml|md|properties))"
)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def slugify(text: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9]+", "-", (text or "").strip().lower()).strip("-")
    return value or "section"


def extract_heading_level(style_name: str) -> int | None:
    if not style_name:
        return None
    match = re.match(r"Heading\s+(\d+)", style_name)
    if not match:
        return None
    return int(match.group(1))


def paragraph_has_image(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def paragraph_image_relationship_ids(paragraph) -> list[str]:
    relation_ids = paragraph._element.xpath(
        ".//*[local-name()='blip']/@*[local-name()='embed']"
    )
    seen: list[str] = []
    for relation_id in relation_ids:
        if relation_id not in seen:
            seen.append(relation_id)
    return seen


def paragraph_font_names(paragraph) -> list[str]:
    names: list[str] = []
    for run in paragraph.runs:
        font_name = run.font.name
        if font_name and font_name not in names:
            names.append(font_name)
    return names


def paragraph_image_extents_pt(paragraph) -> list[dict]:
    extents = []
    seen = set()
    for node in paragraph._element.xpath(".//*[local-name()='extent' and @cx and @cy]"):
        cx = node.get("cx")
        cy = node.get("cy")
        if not cx or not cy:
            continue
        width_pt = round(int(cx) / EMU_PER_PT, 2)
        height_pt = round(int(cy) / EMU_PER_PT, 2)
        key = (width_pt, height_pt)
        if key in seen:
            continue
        seen.add(key)
        extents.append({"width_pt": width_pt, "height_pt": height_pt})
    return extents


def half_points_to_pt(value: str | None) -> float | None:
    if value is None:
        return None
    return round(int(value) / 2.0, 2)


def twips_to_pt(value: str | None) -> float | None:
    if value is None:
        return None
    return round(int(value) / 20.0, 2)


def parse_on_off_value(raw: str | None, *, present: bool) -> bool | None:
    if not present:
        return None
    if raw is None:
        return True
    if raw.lower() in {"0", "false", "off"}:
        return False
    return True


def extract_direct_paragraph_formatting(paragraph) -> dict:
    ppr_nodes = paragraph._element.xpath("./*[local-name()='pPr']")
    if not ppr_nodes:
        return {}

    ppr = ppr_nodes[0]
    alignment = ppr.xpath("./*[local-name()='jc']/@*[local-name()='val']")
    left_indent = ppr.xpath("./*[local-name()='ind']/@*[local-name()='left']")
    right_indent = ppr.xpath("./*[local-name()='ind']/@*[local-name()='right']")
    first_line_indent = ppr.xpath("./*[local-name()='ind']/@*[local-name()='firstLine']")
    first_line_chars = ppr.xpath("./*[local-name()='ind']/@*[local-name()='firstLineChars']")
    space_before = ppr.xpath("./*[local-name()='spacing']/@*[local-name()='before']")
    space_after = ppr.xpath("./*[local-name()='spacing']/@*[local-name()='after']")
    line_spacing = ppr.xpath("./*[local-name()='spacing']/@*[local-name()='line']")
    line_rule = ppr.xpath("./*[local-name()='spacing']/@*[local-name()='lineRule']")

    direct = {}
    if alignment:
        direct["alignment"] = alignment[0].lower()
    if left_indent:
        direct["left_indent_pt"] = twips_to_pt(left_indent[0])
    if right_indent:
        direct["right_indent_pt"] = twips_to_pt(right_indent[0])
    if first_line_indent:
        direct["first_line_indent_pt"] = twips_to_pt(first_line_indent[0])
    if first_line_chars:
        direct["first_line_chars"] = first_line_chars[0]
    if space_before:
        direct["space_before_pt"] = twips_to_pt(space_before[0])
    if space_after:
        direct["space_after_pt"] = twips_to_pt(space_after[0])
    if line_rule:
        direct["line_rule"] = line_rule[0]
    if line_spacing:
        direct["line_spacing_raw"] = line_spacing[0]
        if line_rule and line_rule[0] == "auto":
            direct["line_spacing_multiple"] = round(int(line_spacing[0]) / 240.0, 2)
    return direct


def extract_direct_run_formatting(run) -> dict:
    rpr_nodes = run._element.xpath("./*[local-name()='rPr']")
    if not rpr_nodes:
        return {}

    rpr = rpr_nodes[0]
    ascii_font = rpr.xpath("./*[local-name()='rFonts']/@*[local-name()='ascii']")
    hansi_font = rpr.xpath("./*[local-name()='rFonts']/@*[local-name()='hAnsi']")
    east_asia_font = rpr.xpath("./*[local-name()='rFonts']/@*[local-name()='eastAsia']")
    size = rpr.xpath("./*[local-name()='sz']/@*[local-name()='val']")
    complex_size = rpr.xpath("./*[local-name()='szCs']/@*[local-name()='val']")
    bold_nodes = rpr.xpath("./*[local-name()='b']")
    bold_raw = []
    if bold_nodes:
        bold_raw = bold_nodes[0].xpath("./@*[local-name()='val']")

    direct = {}
    if ascii_font:
        direct["ascii_font"] = ascii_font[0]
    if hansi_font:
        direct["hansi_font"] = hansi_font[0]
    if east_asia_font:
        direct["east_asia_font"] = east_asia_font[0]
    if size:
        direct["size_pt"] = half_points_to_pt(size[0])
    if complex_size:
        direct["complex_size_pt"] = half_points_to_pt(complex_size[0])
    bold = parse_on_off_value(bold_raw[0] if bold_raw else None, present=bool(bold_nodes))
    if bold is not None:
        direct["bold"] = bold
    return direct


def paragraph_fingerprint(
    *,
    index: int,
    style_name: str,
    text: str,
    has_image: bool,
    image_relationship_ids: list[str],
) -> str:
    payload = "||".join(
        [
            str(index),
            style_name or "",
            normalize_text(text),
            "1" if has_image else "0",
            ",".join(image_relationship_ids),
        ]
    )
    return sha1(payload.encode("utf-8")).hexdigest()


def parse_figure_caption(text: str) -> dict | None:
    match = FIGURE_PATTERN.match(normalize_text(text))
    if not match:
        return None
    return {
        "label": match.group(1),
        "figure_id": match.group(2),
        "caption_tail": match.group(3),
    }


def load_docx_document(docx_path: Path):
    return Document(str(docx_path))


def build_image_relationship_inventory(document) -> dict:
    image_relationships = {}
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        image_relationships[rel.rId] = {
            "target_ref": str(rel.target_ref),
            "filename": Path(str(rel.target_ref)).name,
            "reltype": rel.reltype,
        }
    return image_relationships


def build_paragraph_inventory(document) -> tuple[list[dict], list[dict], dict, dict]:
    paragraph_inventory = []
    headings = []
    style_counter: Counter[str] = Counter()
    font_counter: Counter[str] = Counter()
    section_slug_counter: Counter[str] = Counter()

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text or ""
        normalized = normalize_text(text)
        style_name = paragraph.style.name if paragraph.style else ""
        heading_level = extract_heading_level(style_name)
        has_image = paragraph_has_image(paragraph)
        relationship_ids = paragraph_image_relationship_ids(paragraph)
        fonts = paragraph_font_names(paragraph)
        image_extents = paragraph_image_extents_pt(paragraph)
        paragraph_format = paragraph.paragraph_format

        style_counter[style_name or "[no-style]"] += 1
        for font_name in fonts:
            font_counter[font_name] += 1

        fingerprint = paragraph_fingerprint(
            index=index,
            style_name=style_name,
            text=text,
            has_image=has_image,
            image_relationship_ids=relationship_ids,
        )

        metadata = {
            "index": index,
            "style_name": style_name,
            "heading_level": heading_level,
            "text": text,
            "text_preview": normalized[:160],
            "has_image": has_image,
            "image_relationship_ids": relationship_ids,
            "image_extents_pt": image_extents,
            "font_names": fonts,
            "alignment": None
            if paragraph_format.alignment is None
            else paragraph_format.alignment.name.lower(),
            "left_indent_pt": None
            if paragraph_format.left_indent is None
            else round(paragraph_format.left_indent.pt, 2),
            "right_indent_pt": None
            if paragraph_format.right_indent is None
            else round(paragraph_format.right_indent.pt, 2),
            "first_line_indent_pt": None
            if paragraph_format.first_line_indent is None
            else round(paragraph_format.first_line_indent.pt, 2),
            "space_before_pt": None
            if paragraph_format.space_before is None
            else round(paragraph_format.space_before.pt, 2),
            "space_after_pt": None
            if paragraph_format.space_after is None
            else round(paragraph_format.space_after.pt, 2),
            "fingerprint": fingerprint,
        }
        paragraph_inventory.append(metadata)

        if heading_level is not None and normalized:
            base_slug = slugify(normalized)
            section_slug_counter[base_slug] += 1
            count = section_slug_counter[base_slug]
            section_id = base_slug if count == 1 else f"{base_slug}-{count}"
            headings.append(
                {
                    "index": index,
                    "level": heading_level,
                    "text": normalized,
                    "style_name": style_name,
                    "section_id": section_id,
                    "fingerprint": fingerprint,
                }
            )

    return (
        paragraph_inventory,
        headings,
        dict(style_counter.most_common()),
        dict(font_counter.most_common()),
    )


def find_bound_image_paragraph(paragraph_inventory: list[dict], caption_index: int) -> dict | None:
    image_paragraph = None
    for offset in (1, 2, 3):
        candidate_index = caption_index - offset
        if candidate_index < 0:
            continue
        candidate = paragraph_inventory[candidate_index]
        if candidate["has_image"]:
            image_paragraph = candidate
            break

    if image_paragraph is None:
        for offset in (1,):
            candidate_index = caption_index + offset
            if candidate_index >= len(paragraph_inventory):
                continue
            candidate = paragraph_inventory[candidate_index]
            if candidate["has_image"]:
                image_paragraph = candidate
                break

    return image_paragraph


def build_figure_inventory(paragraph_inventory: list[dict]) -> list[dict]:
    figure_inventory = []
    for paragraph in paragraph_inventory:
        caption = parse_figure_caption(paragraph["text"])
        if not caption:
            continue

        image_paragraph = find_bound_image_paragraph(paragraph_inventory, paragraph["index"])

        figure_inventory.append(
            {
                "caption_paragraph_index": paragraph["index"],
                "caption_text": normalize_text(paragraph["text"]),
                "caption_fingerprint": paragraph["fingerprint"],
                "figure_id": caption["figure_id"],
                "figure_label": caption["label"],
                "caption_tail": caption["caption_tail"],
                "image_paragraph_index": None if image_paragraph is None else image_paragraph["index"],
                "image_relationship_ids": []
                if image_paragraph is None
                else image_paragraph["image_relationship_ids"],
                "image_extents_pt": []
                if image_paragraph is None
                else image_paragraph["image_extents_pt"],
            }
        )
    return figure_inventory


def build_section_inventory(document) -> list[dict]:
    sections = []
    for index, section in enumerate(document.sections):
        page_width_pt = round(section.page_width.pt, 2)
        page_height_pt = round(section.page_height.pt, 2)
        top_margin_pt = round(section.top_margin.pt, 2)
        bottom_margin_pt = round(section.bottom_margin.pt, 2)
        left_margin_pt = round(section.left_margin.pt, 2)
        right_margin_pt = round(section.right_margin.pt, 2)
        sections.append(
            {
                "index": index,
                "page_width_pt": page_width_pt,
                "page_height_pt": page_height_pt,
                "top_margin_pt": top_margin_pt,
                "bottom_margin_pt": bottom_margin_pt,
                "left_margin_pt": left_margin_pt,
                "right_margin_pt": right_margin_pt,
                "header_distance_pt": round(section.header_distance.pt, 2),
                "footer_distance_pt": round(section.footer_distance.pt, 2),
                "usable_body_width_pt": round(page_width_pt - left_margin_pt - right_margin_pt, 2),
                "usable_body_height_pt": round(page_height_pt - top_margin_pt - bottom_margin_pt, 2),
            }
        )
    return sections


def build_appendix_inventory(paragraph_inventory: list[dict], headings: list[dict]) -> dict:
    appendix_headings = [
        heading for heading in headings if "appendix" in heading["text"].lower()
    ]
    in_appendix = False
    appendix_file_references: list[dict] = []
    appendix_code_candidates: list[dict] = []

    for paragraph in paragraph_inventory:
        if paragraph["heading_level"] is not None:
            in_appendix = "appendix" in paragraph["text"].lower()
            continue
        if not in_appendix:
            continue

        for match in FILE_REFERENCE_PATTERN.finditer(paragraph["text"]):
            appendix_file_references.append(
                {
                    "paragraph_index": paragraph["index"],
                    "reference": match.group(1),
                    "fingerprint": paragraph["fingerprint"],
                }
            )

        preview = paragraph["text_preview"]
        if len(preview) >= 24 and any(token in preview for token in ("{", "}", ";", "=>", "public ", "const ")):
            appendix_code_candidates.append(
                {
                    "paragraph_index": paragraph["index"],
                    "text_preview": preview,
                    "fingerprint": paragraph["fingerprint"],
                }
            )
    return {
        "headings": appendix_headings,
        "file_references": appendix_file_references,
        "code_candidates": appendix_code_candidates,
    }


def build_caption_image_bindings(figure_inventory: list[dict]) -> list[dict]:
    caption_image_bindings = [
        {
            "figure_id": figure["figure_id"],
            "figure_label": figure["figure_label"],
            "caption_paragraph_index": figure["caption_paragraph_index"],
            "image_paragraph_index": figure["image_paragraph_index"],
            "image_relationship_ids": figure["image_relationship_ids"],
        }
        for figure in figure_inventory
    ]
    return caption_image_bindings


def build_inventory(document, *, source_docx: Path, generated_at_utc: str | None = None) -> dict:
    image_relationships = build_image_relationship_inventory(document)
    paragraph_inventory, headings, paragraph_styles, font_names = build_paragraph_inventory(document)
    figure_inventory = build_figure_inventory(paragraph_inventory)
    appendix_inventory = build_appendix_inventory(paragraph_inventory, headings)
    caption_image_bindings = build_caption_image_bindings(figure_inventory)
    sections = build_section_inventory(document)

    return {
        "source_docx": str(source_docx),
        "generated_at_utc": generated_at_utc or datetime.now(timezone.utc).isoformat(),
        "summary": {
            "paragraph_count": len(paragraph_inventory),
            "heading_count": len(headings),
            "figure_count": len(figure_inventory),
            "image_relationship_count": len(image_relationships),
            "image_paragraph_count": sum(1 for paragraph in paragraph_inventory if paragraph["has_image"]),
            "section_count": len(sections),
            "appendix_heading_count": len(appendix_inventory["headings"]),
            "appendix_file_reference_count": len(appendix_inventory["file_references"]),
        },
        "sections": sections,
        "headings": headings,
        "figures": figure_inventory,
        "caption_image_bindings": caption_image_bindings,
        "image_relationships": image_relationships,
        "paragraph_fingerprints": [
            {
                "index": paragraph["index"],
                "style_name": paragraph["style_name"],
                "text_preview": paragraph["text_preview"],
                "has_image": paragraph["has_image"],
                "image_relationship_ids": paragraph["image_relationship_ids"],
                "image_extents_pt": paragraph["image_extents_pt"],
                "alignment": paragraph["alignment"],
                "left_indent_pt": paragraph["left_indent_pt"],
                "right_indent_pt": paragraph["right_indent_pt"],
                "first_line_indent_pt": paragraph["first_line_indent_pt"],
                "space_before_pt": paragraph["space_before_pt"],
                "space_after_pt": paragraph["space_after_pt"],
                "fingerprint": paragraph["fingerprint"],
            }
            for paragraph in paragraph_inventory
        ],
        "styles": {
            "paragraph_styles": paragraph_styles,
            "font_names": font_names,
        },
        "appendix": appendix_inventory,
    }


def inspect_docx(docx_path: Path) -> dict:
    document = load_docx_document(docx_path)
    return build_inventory(document, source_docx=docx_path)


def build_override_inventory(document) -> dict:
    paragraphs = []
    direct_paragraph_count = 0
    paragraphs_with_run_overrides = 0
    run_override_count = 0

    for index, paragraph in enumerate(document.paragraphs):
        direct_paragraph_formatting = extract_direct_paragraph_formatting(paragraph)
        run_overrides = []
        has_image = paragraph_has_image(paragraph)
        relationship_ids = paragraph_image_relationship_ids(paragraph)

        for run_index, run in enumerate(paragraph.runs):
            direct_run_formatting = extract_direct_run_formatting(run)
            if not direct_run_formatting:
                continue

            run_overrides.append(
                {
                    "run_index": run_index,
                    "text_preview": normalize_text(run.text)[:80],
                    "direct_formatting": direct_run_formatting,
                }
            )

        paragraph_run_override_count = len(run_overrides)

        if not direct_paragraph_formatting and paragraph_run_override_count == 0:
            continue

        if direct_paragraph_formatting:
            direct_paragraph_count += 1
        if paragraph_run_override_count:
            paragraphs_with_run_overrides += 1
            run_override_count += paragraph_run_override_count

        text = paragraph.text or ""
        style_name = paragraph.style.name if paragraph.style else ""
        paragraphs.append(
            {
                "index": index,
                "style_name": style_name,
                "heading_level": extract_heading_level(style_name),
                "text_preview": normalize_text(text)[:160],
                "has_image": has_image,
                "fingerprint": paragraph_fingerprint(
                    index=index,
                    style_name=style_name,
                    text=text,
                    has_image=has_image,
                    image_relationship_ids=relationship_ids,
                ),
                "direct_paragraph_formatting": direct_paragraph_formatting,
                "run_override_count": paragraph_run_override_count,
                "run_overrides": run_overrides,
            }
        )

    return {
        "summary": {
            "paragraphs_with_direct_paragraph_formatting": direct_paragraph_count,
            "paragraphs_with_run_overrides": paragraphs_with_run_overrides,
            "run_override_count": run_override_count,
        },
        "paragraphs": paragraphs,
    }


def inspect_docx_overrides(docx_path: Path) -> dict:
    document = load_docx_document(docx_path)
    return build_override_inventory(document)
