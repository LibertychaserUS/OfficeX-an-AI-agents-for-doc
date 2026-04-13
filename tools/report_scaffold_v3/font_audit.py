from __future__ import annotations

from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .models import FontAuditReport, FontViolationExample, FontViolationGroup


def paragraph_has_only_images(paragraph) -> bool:
    if paragraph.text.strip():
        return False
    drawings = paragraph._element.findall(".//" + qn("w:drawing"))
    return bool(drawings)


def extract_explicit_run_font(run) -> str | None:
    if run.font.name:
        return run.font.name

    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None

    for attr_name in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        value = rfonts.get(qn(attr_name))
        if value:
            return value
    return None


def scan_docx_fonts(docx_path: Path, *, expected_font: str = "Times New Roman") -> FontAuditReport:
    document = Document(str(docx_path))
    grouped_examples: dict[str, list[FontViolationExample]] = defaultdict(list)
    grouped_counts: dict[str, int] = defaultdict(int)

    total_runs = 0
    explicit_expected_runs = 0
    inherited_runs = 0
    explicit_other_runs = 0

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if paragraph_has_only_images(paragraph):
            continue
        snippet = paragraph.text.strip()
        if not snippet:
            continue

        for run in paragraph.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            font_name = extract_explicit_run_font(run)
            if font_name is None:
                inherited_runs += 1
                continue
            if font_name == expected_font:
                explicit_expected_runs += 1
                continue

            explicit_other_runs += 1
            grouped_counts[font_name] += 1
            if len(grouped_examples[font_name]) < 5:
                grouped_examples[font_name].append(
                    FontViolationExample(
                        location=f"Paragraph {paragraph_index}",
                        snippet=snippet[:100],
                    )
                )

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    snippet = paragraph.text.strip()
                    if not snippet:
                        continue

                    for run in paragraph.runs:
                        if not run.text.strip():
                            continue
                        total_runs += 1
                        font_name = extract_explicit_run_font(run)
                        if font_name is None:
                            inherited_runs += 1
                            continue
                        if font_name == expected_font:
                            explicit_expected_runs += 1
                            continue

                        explicit_other_runs += 1
                        grouped_counts[font_name] += 1
                        if len(grouped_examples[font_name]) < 5:
                            grouped_examples[font_name].append(
                                FontViolationExample(
                                    location=f"Table {table_index}, Row {row_index}, Cell {cell_index}",
                                    snippet=snippet[:100],
                                )
                            )

    violations = [
        FontViolationGroup(
            font_name=font_name,
            occurrences=grouped_counts[font_name],
            examples=grouped_examples[font_name],
        )
        for font_name in sorted(grouped_counts)
    ]

    return FontAuditReport(
        source_docx=docx_path,
        expected_font=expected_font,
        total_runs_scanned=total_runs,
        explicit_expected_font_runs=explicit_expected_runs,
        inherited_font_runs=inherited_runs,
        explicit_other_font_runs=explicit_other_runs,
        violations=violations,
    )


def render_font_audit_markdown(report: FontAuditReport) -> str:
    lines = [
        "# Font Audit Report",
        "",
        f"- Source docx: `{report.source_docx}`",
        f"- Expected font: `{report.expected_font}`",
        f"- Total runs scanned: {report.total_runs_scanned}",
        f"- Explicit expected-font runs: {report.explicit_expected_font_runs}",
        f"- Inherited-font runs: {report.inherited_font_runs}",
        f"- Explicit other-font runs: {report.explicit_other_font_runs}",
        "",
        "## Violations",
        "",
    ]

    if not report.violations:
        lines.append("- No explicit non-expected-font violations were found.")
        return "\n".join(lines)

    for violation in report.violations:
        lines.append(f"- `{violation.font_name}`: {violation.occurrences} occurrence(s)")
        for example in violation.examples:
            lines.append(f"  - {example.location}: {example.snippet}")
    return "\n".join(lines)
