from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def normalize_text(text: str) -> str:
    return " ".join(text.split())


def find_paragraph_exact(document: Document, text: str) -> Paragraph:
    target = normalize_text(text)
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text) == target:
            return paragraph
    raise ValueError(f"Could not find paragraph: {text!r}")


def find_paragraph_starting(document: Document, text: str) -> Paragraph:
    target = normalize_text(text)
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text).startswith(target):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with: {text!r}")


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    p = paragraph._p
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    paragraph.add_run(text)


def move_after(anchor: Paragraph, paragraphs: list[Paragraph]) -> Paragraph:
    current = anchor._p
    for paragraph in paragraphs:
        current.addnext(paragraph._p)
        current = paragraph._p
    return Paragraph(current, anchor._parent)


def split_heading_and_body(paragraph: Paragraph, heading_text: str, heading_style: str) -> Paragraph | None:
    full_text = paragraph.text
    if not full_text.startswith(heading_text):
        raise ValueError(f"Paragraph does not start with expected heading: {heading_text!r}")
    body = full_text[len(heading_text):].strip()
    paragraph.style = heading_style
    replace_paragraph_text(paragraph, heading_text)
    if body:
        return insert_paragraph_after(paragraph, body, "Normal")
    return None


def apply_first_review_structure_batch(docx_path: Path) -> None:
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    # REV-001: Separate 1.2.2 from the stray sentence and keep the sentence
    # under the functional-requirements subsection where it belongs.
    p29 = paragraphs[29]
    p30 = paragraphs[30]
    functional_intro = "Functional requirements were derived using user story techniques (Cohn, 2004) to ensure alignment with end-user needs."
    insert_paragraph_after(p29, functional_intro, "Normal")
    p30.style = "Heading 3"
    replace_paragraph_text(p30, "1.2.2 Non-functional Requirements")

    # REV-002: Promote 1.8 to a real heading.
    p107 = paragraphs[107]
    p107.style = "Heading 2"

    # REV-003: Rebind resources and risk content to the correct headings.
    p115 = paragraphs[115]
    p116 = paragraphs[116]
    p117 = paragraphs[117]
    p118 = paragraphs[118]
    p119 = paragraphs[119]
    p120 = paragraphs[120]
    p121 = paragraphs[121]
    p122 = paragraphs[122]
    p123 = paragraphs[123]
    p122.style = "Heading 2"
    move_after(p115, [p123, p122, p116, p117, p118, p119, p120, p121])

    # REV-004: Promote 2.7.2 to a real subheading.
    p15428 = paragraphs[15428]
    p15428.style = "Heading 3"

    # REV-005: Split 2.8.1 and 2.8.2 into clean subheadings.
    split_heading_and_body(paragraphs[15433], "2.8.1 Installation Guide", "Heading 3")
    split_heading_and_body(paragraphs[15435], "2.8.2 User Manual", "Heading 3")

    # REV-006: Separate Stage 3 heading paragraphs from their embedded body text.
    split_heading_and_body(
        paragraphs[15459],
        "3.2 Evaluate the Strengths and Weaknesses of the Practical Task Results",
        "Heading 2",
    )
    split_heading_and_body(
        paragraphs[15470],
        "3.3 Recommendations for Future Development",
        "Heading 2",
    )
    split_heading_and_body(
        paragraphs[15479],
        "3.4 Summary of Modifications Made During the Project",
        "Heading 2",
    )
    split_heading_and_body(
        paragraphs[15487],
        "3.5 Knowledge and Skills Gained",
        "Heading 2",
    )

    doc.save(docx_path)


def apply_first_review_content_batch(docx_path: Path) -> None:
    doc = Document(docx_path)

    # REV-007: Strengthen requirement-to-implementation traceability in 1.2.
    requirement_intro = find_paragraph_exact(
        doc,
        "Functional requirements were derived using user story techniques (Cohn, 2004) to ensure alignment with end-user needs.",
    )
    insert_paragraph_after(
        requirement_intro,
        (
            "To keep the requirements actionable during implementation, I grouped the functional requirements into "
            "coherent delivery slices rather than treating Table 1 as an isolated checklist: identity and profile "
            "(FR-01 to FR-04), discovery and listing lifecycle (FR-05 to FR-12), engagement features "
            "(FR-13 to FR-15), and governance plus platform-quality features (FR-16 to FR-23). This grouping is "
            "carried forward into the domain analysis in Section 1.12, the design artefacts in Section 1.13, the "
            "backend and UI code walkthroughs in Sections 2.1 and 2.2, and the requirement-by-requirement "
            "assessment in Section 3.1."
        ),
        "Normal",
    )

    # REV-008: Add explicit user-analysis and data-binding signposting in 1.12.
    system_analysis = find_paragraph_exact(doc, "Business Model Analysis")
    system_analysis = insert_paragraph_after(
        system_analysis,
        (
            "This analysis was carried out from both a user perspective and a data-binding perspective. From the "
            "user perspective, the four actors identified in the use case model drive the structure of the system: "
            "guests browse discovery pages, buyers place orders and initiate inquiries, sellers manage listings and "
            "fulfil transactions, and administrators govern users, reports, announcements, and audit trails. That "
            "actor analysis explains why the later design separates public routes, buyer workflows, seller workspace "
            "screens, and admin governance surfaces instead of presenting one undifferentiated interface."
        ),
        "Normal",
    )
    insert_paragraph_after(
        system_analysis,
        (
            "From the data-binding perspective, the same actions are anchored to explicit domain relationships. "
            "Listings bind to categories and ordered media; orders bind buyers, sellers, and listing state changes; "
            "conversations and notifications bind communication back to a specific listing or transaction; and "
            "moderation records bind reports and administrative actions back to the affected content. Those "
            "relationships are then preserved in the class diagrams in Section 1.13 and in the module slices "
            "described in Section 2.1."
        ),
        "Normal",
    )

    # REV-009: Strengthen design-to-code traceability in 2.1.
    backend_modules = find_paragraph_starting(
        doc,
        "The backend source code is organised into the following domain packages:",
    )
    insert_paragraph_after(
        backend_modules,
        (
            "This package breakdown is a direct implementation of the domain analysis rather than an arbitrary folder "
            "structure. UserAccount and UserProfile from Section 1.12 map to the identity and profile packages; "
            "Listing, Category, and ListingMedia map to catalog, listing, and media; Order and Review map to "
            "orders and review; Conversation, Message, Favorite, Notification, and ListingReport map to inquiry, "
            "favorite, notification, and moderation; and platform-governance concerns map to admin and dashboard. "
            "The same slice boundaries are exercised by backend tests such as AuthServiceTest, "
            "ListingMediaServiceTest, ListingFlowIntegrationTest, OrderFlowIntegrationTest, "
            "InquiryFlowIntegrationTest, FavoritesDashboardIntegrationTest, CommentsReadModelsIntegrationTest, and "
            "AdminModerationIntegrationTest, so the code listing can be read as traceable implementation of the "
            "earlier design rather than as isolated source files."
        ),
        "Normal",
    )

    # REV-010: Strengthen wireframe-to-route/component linkage in 2.2.
    frontend_intro = find_paragraph_starting(
        doc,
        "This section presents the frontend implementation of the LoopMart application.",
    )
    insert_paragraph_after(
        frontend_intro,
        (
            "The UI code is organised to mirror the wireframe families and user journeys introduced earlier in the "
            "design sections. App.jsx groups public discovery routes (home, search, categories, listing detail, "
            "public profiles), buyer routes (orders, favorites, buyer dashboard), seller workspace routes "
            "(inventory, editor, earnings, reviews), and admin governance routes (users, listings, reports, "
            "categories, orders, reviews, audit log, announcements, settings). Within those routes, components such "
            "as AuthPage, ListingEditorPage, ProfilePage, OrdersPage, and SearchResultsPage demonstrate how the "
            "planned wireframes were translated into working screens with required fields, route-state handling, API "
            "integration, and visible feedback panels."
        ),
        "Normal",
    )

    # REV-011: Clarify unfamiliar-library learning process and rationale.
    springdoc_paragraph = find_paragraph_starting(
        doc,
        "SpringDoc OpenAPI: Automatically generates OpenAPI 3.0 documentation",
    )
    insert_paragraph_after(
        springdoc_paragraph,
        (
            "The choice of these libraries was based on risk reduction as much as convenience. I considered staying "
            "closer to the taught baseline with more manual servlet-style wiring, hand-managed SQL changes, and a "
            "separately maintained API reference, but that would have increased repetitive infrastructure work and "
            "weakened reproducibility. The learning process therefore followed a staged pattern: I read the official "
            "documentation, built a small proof-of-concept for one slice, and only then expanded the library into "
            "the wider system. This was especially important for Spring Boot, Flyway, React Router, and SpringDoc "
            "because each introduced framework conventions that were new to me."
        ),
        "Normal",
    )

    # REV-012: Document the shared backend/frontend error contract.
    error_handling = find_paragraph_starting(
        doc,
        "The application implements a centralised error handling strategy on the backend",
    )
    insert_paragraph_after(
        error_handling,
        (
            "On the frontend, the same contract is preserved by the shared api/client.js module and the page-level "
            "error panels. parseResponse converts backend error payloads into JavaScript exceptions, handle401 "
            "redirects unauthenticated users back to /login, and pages such as AuthPage, ListingEditorPage, "
            "OrdersPage, and ProfilePage render the returned messages through consistent message-panel-error "
            "components. As a result, validation and exception handling are not confined to the business model; they "
            "are surfaced to users in a predictable way across both backend and UI layers."
        ),
        "Normal",
    )

    # REV-013: Rebalance internal-documentation claims more honestly.
    internal_docs = find_paragraph_starting(
        doc,
        "The project maintains internal documentation through several mechanisms:",
    )
    insert_paragraph_after(
        internal_docs,
        (
            "The code-comment density inside src/main is lighter than in a tutorial-style codebase, so I do not "
            "claim that every class is exhaustively annotated inline. Instead, the stronger internal-documentation "
            "trail comes from the combination of descriptive package and controller names, migration history, test "
            "classes, README and DEMO_GUIDE assets, the start.sh launcher, the user manual, and the project trace "
            "and checkpoint records. This is a more accurate description of the evidence base: the project is "
            "documented, but that documentation is distributed across code structure and project artefacts rather "
            "than concentrated only in comments."
        ),
        "Normal",
    )

    # REV-014: Separate automated backend evidence from manual UI evidence.
    test_strategy = find_paragraph_starting(
        doc,
        "The test strategy employs a multi-layered approach combining unit testing",
    )
    insert_paragraph_after(
        test_strategy,
        (
            "The evidence was intentionally separated into three categories. First, automated unit tests check "
            "service-level rules such as registration/login rejection and media constraints (for example "
            "AuthServiceTest and ListingMediaServiceTest). Second, automated integration tests use MockMvc against "
            "the running Spring context to validate end-to-end backend flows such as listing publication, order "
            "lifecycle, inquiry reuse, favorites/dashboard metrics, comments read models, and admin moderation. "
            "Third, manual browser walkthrough evidence validates the visible UI workflow. This separation matters "
            "because a screenshot can demonstrate a successful screen state, but it cannot substitute for automated "
            "assertions on business rules or API contracts."
        ),
        "Normal",
    )
    test_runs_intro = find_paragraph_exact(
        doc,
        "The following evidence demonstrates that the test cases identified in the test plan have been executed successfully.",
    )
    insert_paragraph_after(
        test_runs_intro,
        (
            "Accordingly, the evidence below should be read in two streams. The automated stream is represented by "
            "JUnit and MockMvc outputs for AuthServiceTest, ListingFlowIntegrationTest, OrderFlowIntegrationTest, "
            "InquiryFlowIntegrationTest, FavoritesDashboardIntegrationTest, CommentsReadModelsIntegrationTest, and "
            "AdminModerationIntegrationTest. The manual stream is represented by Figures 48 to 60, which capture the "
            "visible success state of key user journeys. Together they show both internal correctness and external "
            "usability."
        ),
        "Normal",
    )

    # REV-015: Align user-documentation claims to the real startup/help assets.
    installation_paragraph = find_paragraph_starting(
        doc,
        "Prerequisites: Java 21",
    )
    replace_paragraph_text(
        installation_paragraph,
        (
            "Prerequisites: Java 21+, Apache Maven 3.8+, Node.js 18+, and a running PostgreSQL instance for the "
            "loopmart_rebuild database. The recommended startup path is the project-root start.sh script, which "
            "checks dependencies, clears ports 8080 and 5173, starts the Spring Boot backend and Vite frontend, and "
            "prints the health-check and Swagger UI URLs. A manual two-terminal startup path is also supported by "
            "running Maven from apps/backend and npm run dev from apps/frontend."
        ),
    )
    user_manual_intro = find_paragraph_exact(
        doc,
        "This section provides guidance for end users interacting with the LoopMart application.",
    )
    insert_paragraph_after(
        user_manual_intro,
        (
            "The report summary here is aligned to the real support artefacts shipped with the project: the root "
            "start.sh launcher, the bilingual USER_MANUAL.md file, and the /help route inside the application. I am "
            "therefore using the report to summarise those assets rather than inventing a separate set of "
            "instructions that diverges from the running system."
        ),
        "Normal",
    )
    registration_paragraph = find_paragraph_starting(
        doc,
        "Registration and Login: Navigate to the authentication page by clicking",
    )
    replace_paragraph_text(
        registration_paragraph,
        (
            "Registration and Login: Navigate to the authentication page by clicking 'Sign in' or 'Join free' in "
            "the navigation bar. New users provide a username, display name, password, and the current math CAPTCHA "
            "answer; returning users provide username, password, and CAPTCHA verification. Password recovery is "
            "handled through the dedicated forgot-password and reset-password routes, matching the documented flow "
            "in the project user manual."
        ),
    )

    # REV-016: Make the requirement-mapping claim more visibly traceable.
    traceability_intro = find_paragraph_starting(
        doc,
        "To provide a structured assessment, I mapped each of the 23 functional requirements",
    )
    insert_paragraph_after(
        traceability_intro,
        (
            "The traceability table should be read as the integration point between planning, implementation, and "
            "validation. Each row links the requirement list from Section 1.2 to a concrete implementation "
            "entrypoint and then to a named test case or UI evidence item. In practice, the rows cluster into "
            "identity/profile, listing/catalog/media, orders/reviews, inquiry/notification/favorites, and "
            "governance slices, which matches the module boundaries described in Section 2.1 and the route groups "
            "described in Section 2.2. This makes the satisfaction claim visible rather than purely narrative."
        ),
        "Normal",
    )

    # Also correct one inaccurate evidence label inside the existing FR traceability table.
    for table in doc.tables:
        header = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if header == ["FR ID", "Requirement", "Status", "Evidence", "Notes"]:
            for row in table.rows[1:]:
                if normalize_text(row.cells[0].text) == "FR-08":
                    row.cells[3].text = "TC-18, SearchResultsPage + ListingController"
                    break
            break

    # REV-017: Deepen evaluation logic, prioritisation, and justification in Stage 3.
    strengths_intro = find_paragraph_starting(
        doc,
        "From a technical standpoint, several specific strengths stand out.",
    )
    insert_paragraph_after(
        strengths_intro,
        (
            "To evaluate these results more rigorously, I ranked the outcomes by impact on delivery quality rather "
            "than by novelty alone. The highest-value strengths are the modular backend structure, reproducible "
            "Flyway-backed database evolution, and the presence of real governance surfaces instead of placeholder "
            "admin screens, because these decisions improved correctness, maintainability, and assessment coverage "
            "across the whole project. The highest-priority weaknesses are the lack of real-time communication, "
            "limited client-side validation, and the absence of browser-level end-to-end automation, because these "
            "are the gaps most likely to affect user experience and regression risk."
        ),
        "Normal",
    )
    recommendation_intro = find_paragraph_exact(
        doc,
        "Based on the evaluation of the current prototype, the following recommendations are proposed for future development iterations:",
    )
    insert_paragraph_after(
        recommendation_intro,
        (
            "I would prioritise these recommendations in three tiers. First priority is automated regression coverage "
            "and real-time communication because they directly reduce user-facing defects and address the most "
            "significant weaknesses identified in Section 3.2. Second priority is media optimisation and search "
            "quality because they improve performance and discovery as the marketplace grows. Lower-priority items "
            "such as native mobile apps and broader internationalisation remain worthwhile, but they sit outside the "
            "minimum baseline until the existing web workflow is more fully hardened."
        ),
        "Normal",
    )
    modifications_intro = find_paragraph_starting(
        doc,
        "Several modifications were made to the original project plan during development",
    )
    insert_paragraph_after(
        modifications_intro,
        (
            "These modifications were not random additions. Architecture changes were driven by maintainability and "
            "reproducibility, security changes by identified abuse risk, and UI changes by observed usability "
            "friction during testing. Framing the modifications this way is important because it shows the project "
            "evolved against a baseline plan rather than through untracked feature creep."
        ),
        "Normal",
    )
    learning_intro = find_paragraph_exact(
        doc,
        "This project provided significant opportunities for learning and skill development across multiple technical and professional domains.",
    )
    insert_paragraph_after(
        learning_intro,
        (
            "The most valuable learning did not come from isolated syntax practice, but from understanding cause "
            "and effect between design decisions and delivered outcomes. Choosing a modular architecture improved "
            "testing and traceability; adopting Flyway changed the way I approached database evolution; and "
            "building the authentication, listing, order, and governance flows end to end made me more aware of how "
            "security, validation, and documentation must work together. These are the areas I would now develop "
            "further in future projects, particularly automated testing, real-time communication, and stronger "
            "upfront interface validation."
        ),
        "Normal",
    )

    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply verified first-review fixes to a candidate GU2 report.")
    parser.add_argument("--docx", required=True, help="Path to the candidate docx file to modify in place.")
    parser.add_argument(
        "--batch",
        choices=["first-structure", "first-content", "all"],
        default="all",
        help="Which verified revision batch to apply.",
    )
    args = parser.parse_args()
    docx_path = Path(args.docx)

    if args.batch in {"first-structure", "all"}:
        apply_first_review_structure_batch(docx_path)
    if args.batch in {"first-content", "all"}:
        apply_first_review_content_batch(docx_path)


if __name__ == "__main__":
    main()
