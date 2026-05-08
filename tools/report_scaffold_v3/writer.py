from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from PIL import Image

from .models import (
    BuildResult,
    BuildSourceManifest,
    CodeBlockSpec,
    ImageBlockSpec,
    ImageRoleManifest,
    ParagraphBlockSpec,
    TableBlockSpec,
    TableRoleManifest,
    WriteContractManifest,
    WriteRoleManifest,
)

logger = logging.getLogger(__name__)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class WriteContractError(ValueError):
    pass


KNOWN_PARAGRAPH_FORMAT_FIELDS = frozenset({
    "alignment",
    "first_line_indent_pt",
    "left_indent_pt",
    "right_indent_pt",
    "line_spacing_multiple",
    "space_before_pt",
    "space_after_pt",
})

KNOWN_RUN_FORMAT_FIELDS = frozenset({
    "font_name",
    "size_pt",
    "bold",
    "italic",
})


def _clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _apply_paragraph_format(paragraph, role: WriteRoleManifest) -> None:
    unknown_fields = set(role.paragraph_format) - KNOWN_PARAGRAPH_FORMAT_FIELDS
    if unknown_fields:
        logger.warning(
            "Write contract role `%s` has unknown paragraph_format fields: %s",
            role.style,
            ", ".join(sorted(unknown_fields)),
        )
    paragraph.style = role.style
    paragraph_format = paragraph.paragraph_format
    alignment = role.paragraph_format.get("alignment")
    if alignment:
        paragraph.alignment = ALIGNMENTS[alignment]

    first_line_indent_pt = role.paragraph_format.get("first_line_indent_pt")
    if first_line_indent_pt is not None:
        paragraph_format.first_line_indent = Pt(first_line_indent_pt)

    left_indent_pt = role.paragraph_format.get("left_indent_pt")
    if left_indent_pt is not None:
        paragraph_format.left_indent = Pt(left_indent_pt)

    right_indent_pt = role.paragraph_format.get("right_indent_pt")
    if right_indent_pt is not None:
        paragraph_format.right_indent = Pt(right_indent_pt)

    line_spacing_multiple = role.paragraph_format.get("line_spacing_multiple")
    if line_spacing_multiple is not None:
        paragraph_format.line_spacing = float(line_spacing_multiple)

    space_before_pt = role.paragraph_format.get("space_before_pt")
    if space_before_pt is not None:
        paragraph_format.space_before = Pt(space_before_pt)

    space_after_pt = role.paragraph_format.get("space_after_pt")
    if space_after_pt is not None:
        paragraph_format.space_after = Pt(space_after_pt)


def _apply_run_format(run, role: WriteRoleManifest) -> None:
    run_format = role.run_format
    unknown_fields = set(run_format) - KNOWN_RUN_FORMAT_FIELDS
    if unknown_fields:
        logger.warning(
            "Write contract role `%s` has unknown run_format fields: %s",
            role.style,
            ", ".join(sorted(unknown_fields)),
        )
    font_name = run_format.get("font_name")
    if font_name:
        run.font.name = font_name

    size_pt = run_format.get("size_pt")
    if size_pt is not None:
        run.font.size = Pt(size_pt)

    bold = run_format.get("bold")
    if bold is not None:
        run.bold = bool(bold)

    italic = run_format.get("italic")
    if italic is not None:
        run.italic = bool(italic)


def _natural_image_width_pt(image_path: Path) -> float:
    with Image.open(image_path) as image:
        dpi_x = image.info.get("dpi", (72, 72))[0] or 72
        return image.size[0] * 72.0 / float(dpi_x)


def _resolve_image_width(image_path: Path, role: ImageRoleManifest) -> Pt | None:
    if role.width_mode == "fixed_width_pt" and role.fixed_width_pt is not None:
        return Pt(role.fixed_width_pt)

    if role.width_mode == "fit_usable_body_width" and role.max_width_pt is not None:
        natural_width_pt = _natural_image_width_pt(image_path)
        return Pt(min(natural_width_pt, role.max_width_pt))

    return None


def _validate_contract_role(contract: WriteContractManifest, role_name: str) -> WriteRoleManifest:
    role = contract.paragraph_roles.get(role_name)
    if role is None:
        raise WriteContractError(f"Unknown paragraph role `{role_name}`.")
    return role


def _validate_image_role(contract: WriteContractManifest, role_name: str) -> ImageRoleManifest:
    role = contract.image_roles.get(role_name)
    if role is None:
        raise WriteContractError(f"Unknown image role `{role_name}`.")
    return role


def _validate_table_role(contract: WriteContractManifest, role_name: str) -> TableRoleManifest:
    role = contract.table_roles.get(role_name)
    if role is None:
        # Return sensible default if no explicit table role defined
        return TableRoleManifest()
    return role


def _render_paragraph(document: Document, block: ParagraphBlockSpec, role: WriteRoleManifest) -> None:
    paragraph = document.add_paragraph()
    _apply_paragraph_format(paragraph, role)
    run = paragraph.add_run(block.text)
    _apply_run_format(run, role)


def _render_image(
    document: Document,
    block: ImageBlockSpec,
    image_role: ImageRoleManifest,
    caption_role: WriteRoleManifest,
) -> None:
    width = _resolve_image_width(block.image_path, image_role)
    document.add_picture(str(block.image_path), width=width)
    image_paragraph = document.paragraphs[-1]
    if image_role.center_paragraph:
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = document.add_paragraph()
    _apply_paragraph_format(caption, caption_role)
    run = caption.add_run(block.caption)
    _apply_run_format(run, caption_role)


def _render_table(
    document: Document,
    block: TableBlockSpec,
    table_role: TableRoleManifest,
    caption_role: WriteRoleManifest | None,
) -> None:
    """Render a table block into the document."""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Emu

    rows_count = len(block.rows) + (1 if block.headers else 0)
    cols_count = max(
        len(block.headers) if block.headers else 0,
        max((len(row) for row in block.rows), default=0),
    )
    if cols_count == 0 or rows_count == 0:
        return

    table = document.add_table(rows=rows_count, cols=cols_count)
    table.style = table_role.style

    # Apply cell margins via table properties
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = table._tbl._new_tblPr()
    cell_margin = tbl_pr.makeelement(qn("w:tblCellMar"), {})
    for side, dxa_val in [
        ("top", table_role.cell_margin_top_dxa),
        ("bottom", table_role.cell_margin_bottom_dxa),
        ("start", table_role.cell_margin_left_dxa),
        ("end", table_role.cell_margin_right_dxa),
    ]:
        el = cell_margin.makeelement(qn(f"w:{side}"), {
            qn("w:w"): str(dxa_val),
            qn("w:type"): "dxa",
        })
        cell_margin.append(el)
    tbl_pr.append(cell_margin)

    # Fill header row
    row_offset = 0
    if block.headers:
        header_row = table.rows[0]
        for col_idx, header_text in enumerate(block.headers):
            if col_idx < cols_count:
                cell = header_row.cells[col_idx]
                cell.text = header_text
                # Bold header text
                if table_role.header_bold:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                # Header background color
                if table_role.header_bg_color:
                    shading = cell._element.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): table_role.header_bg_color,
                    })
                    cell._element.get_or_add_tcPr().append(shading)
        # Mark header row to repeat across pages
        if table_role.repeat_header_row:
            tr_pr = header_row._tr.get_or_add_trPr()
            tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"), {}))
        row_offset = 1

    # Fill data rows
    for row_idx, row_data in enumerate(block.rows):
        table_row = table.rows[row_idx + row_offset]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < cols_count:
                table_row.cells[col_idx].text = cell_text

    # Apply column widths if specified
    if block.column_widths and len(block.column_widths) == cols_count:
        total_weight = sum(block.column_widths)
        # Use 9360 DXA as default full width (US Letter - 1in margins)
        # This will be overridden by actual page width in practice
        total_dxa = 9360
        for col_idx, weight in enumerate(block.column_widths):
            col_dxa = int(total_dxa * weight / total_weight)
            for row in table.rows:
                cell = row.cells[col_idx]
                cell.width = Emu(col_dxa * 914)  # DXA to EMU

    # Caption
    if block.caption and caption_role:
        caption_para = document.add_paragraph()
        _apply_paragraph_format(caption_para, caption_role)
        run = caption_para.add_run(block.caption)
        _apply_run_format(run, caption_role)


def _execute_code_block(document: Document, block: CodeBlockSpec) -> None:
    """Execute AI-generated python-docx code in a restricted scope.

    The code receives `document` as a python-docx Document object.
    Only python-docx, standard library math/datetime, and docx submodules
    are available. No file I/O, no network, no imports beyond the allowlist.
    """
    if not block.code.strip():
        return

    FORBIDDEN = ["import os", "import sys", "import subprocess", "open(",
                 "__import__", "eval(", "exec(", "compile(", "globals(",
                 "locals(", "getattr(", "setattr(", "delattr(",
                 "shutil", "pathlib", "socket", "http", "urllib"]

    for forbidden in FORBIDDEN:
        if forbidden in block.code:
            logger.warning(
                "Code block rejected: contains forbidden pattern '%s'",
                forbidden,
            )
            return

    # Restricted execution scope
    import docx as docx_module
    from docx.shared import Pt, Emu, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    safe_globals = {
        "__builtins__": {"range": range, "len": len, "str": str,
                         "int": int, "float": float, "bool": bool,
                         "list": list, "dict": dict, "tuple": tuple,
                         "enumerate": enumerate, "zip": zip, "round": round,
                         "min": min, "max": max, "sum": sum, "sorted": sorted,
                         "True": True, "False": False, "None": None},
        "document": document,
        "docx": docx_module,
        "Pt": Pt, "Emu": Emu, "Inches": Inches, "Cm": Cm,
        "RGBColor": RGBColor,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
    }

    try:
        exec(block.code, safe_globals)  # noqa: S102
        logger.debug("Code block executed: %s", block.description or "(unnamed)")
    except Exception as exc:
        logger.warning("Code block execution failed: %s — %s", block.description, exc)


def parse_block_specs(source: BuildSourceManifest) -> list[ParagraphBlockSpec | ImageBlockSpec | TableBlockSpec | CodeBlockSpec]:
    parsed_blocks: list[ParagraphBlockSpec | ImageBlockSpec | TableBlockSpec | CodeBlockSpec] = []
    for block in source.blocks:
        kind = block.get("kind")
        if kind == "paragraph":
            parsed_blocks.append(ParagraphBlockSpec.model_validate(block))
            continue
        if kind == "image":
            parsed_blocks.append(ImageBlockSpec.model_validate(block))
            continue
        if kind == "table":
            parsed_blocks.append(TableBlockSpec.model_validate(block))
            continue
        if kind == "code":
            parsed_blocks.append(CodeBlockSpec.model_validate(block))
            continue
        raise WriteContractError(f"Unsupported block kind `{kind}` in build source.")
    return parsed_blocks


def build_word_candidate(
    *,
    template_docx: Path,
    source: BuildSourceManifest,
    contract: WriteContractManifest,
    output_docx: Path,
) -> BuildResult:
    document = Document(str(template_docx))
    if contract.default_output_strategy.get("clear_existing_body_content", True):
        _clear_document_body(document)

    parsed_blocks = parse_block_specs(source)
    paragraph_count = 0
    image_count = 0
    table_count = 0

    for block in parsed_blocks:
        if isinstance(block, ParagraphBlockSpec):
            role = _validate_contract_role(contract, block.role)
            _render_paragraph(document, block, role)
            paragraph_count += 1
            continue

        if isinstance(block, ImageBlockSpec):
            image_role = _validate_image_role(contract, block.role)
            caption_role = _validate_contract_role(contract, image_role.caption_role)
            _render_image(document, block, image_role, caption_role)
            paragraph_count += 2
            image_count += 1
            continue

        if isinstance(block, TableBlockSpec):
            table_role = _validate_table_role(contract, block.role)
            caption_role_obj = None
            if block.caption and table_role.caption_role:
                caption_role_obj = contract.paragraph_roles.get(table_role.caption_role)
            _render_table(document, block, table_role, caption_role_obj)
            table_count += 1
            if block.caption:
                paragraph_count += 1  # caption paragraph
            continue

        if isinstance(block, CodeBlockSpec):
            _execute_code_block(document, block)
            continue

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))

    return BuildResult(
        document_id=source.document_id,
        template_docx=template_docx,
        output_docx=output_docx,
        block_count=len(parsed_blocks),
        paragraph_count=paragraph_count,
        image_count=image_count,
    )
