"""OfficeX edit runtime: AI-driven editing of existing documents.

Loads an existing .docx, sends its content to AI with user instructions,
parses the AI response as patch operations, and applies them through
the deterministic execution engine.

This completes the create/edit/validate lifecycle.
"""

from __future__ import annotations

import json
import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from .docx_inspector import inspect_docx
from .generate_runtime import _extract_json_from_response, GenerateReport
from .runtime_common import make_local_runtime_identifier, sanitize_runtime_identifier
from .paths import SANDBOXES_DIR

logger = logging.getLogger(__name__)

EDIT_SYSTEM_PROMPT = """You are editing an existing Word document. The user will provide:
1. The current document content (paragraphs with their styles)
2. An edit instruction describing what to change

Return ONLY a JSON object with this structure:
{
  "operations": [
    {
      "type": "replace",
      "paragraph_index": 3,
      "new_text": "Replacement paragraph text"
    },
    {
      "type": "insert_after",
      "paragraph_index": 5,
      "new_text": "New paragraph to insert after paragraph 5",
      "style": "Normal"
    },
    {
      "type": "delete",
      "paragraph_index": 7
    }
  ],
  "summary": "Brief description of what was changed and why"
}

Rules:
- paragraph_index is 0-based, matching the document content provided
- "replace" changes the text of an existing paragraph (preserves style)
- "insert_after" adds a new paragraph after the specified index
- "delete" removes a paragraph
- Only modify what the user asked for. Preserve everything else.
- Return valid JSON only, no markdown fences
"""


def _extract_document_content(docx_path: Path) -> list[dict]:
    """Read document paragraphs for AI context."""
    from docx import Document

    doc = Document(str(docx_path))
    content = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            content.append({
                "index": i,
                "style": para.style.name,
                "text": text[:500],  # Limit to avoid token explosion
            })
    return content


def _apply_edit_operations(
    docx_path: Path,
    output_path: Path,
    operations: list[dict],
) -> dict:
    """Apply edit operations to a docx file."""
    import shutil
    from docx import Document

    # Work on a copy (sandbox principle)
    shutil.copy2(str(docx_path), str(output_path))
    doc = Document(str(output_path))
    paragraphs = list(doc.paragraphs)

    applied = 0
    skipped = 0
    errors = []

    # Sort operations by index descending so deletions/insertions
    # don't shift indices of subsequent operations
    sorted_ops = sorted(operations, key=lambda op: op.get("paragraph_index", 0), reverse=True)

    for op in sorted_ops:
        op_type = op.get("type", "")
        idx = op.get("paragraph_index", -1)

        if idx < 0 or idx >= len(paragraphs):
            errors.append(f"Invalid paragraph_index {idx} (document has {len(paragraphs)} paragraphs)")
            skipped += 1
            continue

        if op_type == "replace":
            new_text = op.get("new_text", "")
            if new_text:
                # Clear existing runs and add new text
                para = paragraphs[idx]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
                applied += 1
            else:
                skipped += 1

        elif op_type == "insert_after":
            new_text = op.get("new_text", "")
            style = op.get("style", "Normal")
            if new_text:
                # Insert after the target paragraph
                new_para = doc.add_paragraph(new_text, style=style)
                # Move to correct position
                target = paragraphs[idx]._element
                target.addnext(new_para._element)
                applied += 1
            else:
                skipped += 1

        elif op_type == "delete":
            para = paragraphs[idx]
            para._element.getparent().remove(para._element)
            applied += 1

        else:
            errors.append(f"Unknown operation type: {op_type}")
            skipped += 1

    doc.save(str(output_path))

    return {
        "applied": applied,
        "skipped": skipped,
        "errors": errors,
    }


def run_edit(
    *,
    docx_path: Path,
    instruction: str,
    provider_id: str = "compatible_local",
    model_id: str | None = None,
    run_id: str | None = None,
    output_dir: Path | None = None,
    include_visual_audit: bool = True,
) -> GenerateReport:
    """Edit an existing document based on user instruction."""
    run_id = run_id or sanitize_runtime_identifier(
        make_local_runtime_identifier("edit"),
        fallback="officex-edit",
    )
    output_dir = (output_dir or SANDBOXES_DIR / run_id).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = docx_path.expanduser().resolve()
    output_docx = output_dir / f"{run_id}.docx"

    # Step 1: Read current document content
    logger.debug("Step 1: Reading document content")
    content = _extract_document_content(docx_path)

    content_text = "\n".join(
        f"[{p['index']}] ({p['style']}) {p['text']}"
        for p in content
    )

    # Step 2: Call AI for edit operations
    logger.debug("Step 2: Calling AI for edit operations")
    api_key = os.environ.get("OFFICEX_PROVIDER_API_KEY")
    if not api_key:
        return GenerateReport(
            run_id=run_id, status="ai_failed",
            error_message="No API key. Set OFFICEX_PROVIDER_API_KEY.",
        )

    base_url = os.environ.get("OFFICEX_PROVIDER_BASE_URL")
    from openai import OpenAI
    client_kwargs: dict = {"api_key": api_key}
    if base_url:
        client_kwargs["base_url"] = base_url
    client = OpenAI(**client_kwargs)
    resolved_model = model_id or "qwen-plus"

    user_prompt = (
        f"Current document content:\n\n{content_text}\n\n"
        f"Edit instruction: {instruction}"
    )

    try:
        response = client.chat.completions.create(
            model=resolved_model,
            messages=[
                {"role": "system", "content": EDIT_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        ai_text = response.choices[0].message.content or ""
        ai_tokens = None
        if response.usage:
            ai_tokens = {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens,
            }
    except Exception as exc:
        return GenerateReport(
            run_id=run_id, status="ai_failed",
            error_message=str(exc),
        )

    (output_dir / "ai_response.txt").write_text(ai_text, encoding="utf-8")

    # Step 3: Parse operations
    logger.debug("Step 3: Parsing edit operations")
    try:
        extracted = _extract_json_from_response(ai_text)
        data = json.loads(extracted)
        operations = data.get("operations", [])
        summary = data.get("summary", "")
    except (json.JSONDecodeError, KeyError) as exc:
        return GenerateReport(
            run_id=run_id, status="ai_failed",
            ai_model=resolved_model, ai_tokens=ai_tokens,
            error_message=f"Failed to parse edit operations: {exc}",
        )

    (output_dir / "operations.json").write_text(
        json.dumps({"operations": operations, "summary": summary}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # Step 4: Apply operations
    logger.debug("Step 4: Applying edit operations")
    try:
        apply_result = _apply_edit_operations(docx_path, output_docx, operations)
    except Exception as exc:
        return GenerateReport(
            run_id=run_id, status="build_failed",
            ai_model=resolved_model, ai_tokens=ai_tokens,
            error_message=f"Failed to apply edits: {exc}",
        )

    (output_dir / "apply_result.json").write_text(
        json.dumps(apply_result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # Step 5: Validation
    logger.debug("Step 5: Structural validation")
    from .manifest_loader import load_layout_contract, load_template_profile
    from .ooxml_inspector import extract_effective_style_inventory
    from .docx_inspector import inspect_docx_overrides
    from .validation import build_validation_report

    template_profile = load_template_profile().model_dump(mode="json")
    layout_contract = load_layout_contract().model_dump(mode="json")
    inventory = inspect_docx(output_docx)
    val_report = build_validation_report(
        output_docx, inventory,
        target_role="candidate_output",
        template_profile=template_profile,
        layout_contract=layout_contract,
        style_inventory=extract_effective_style_inventory(output_docx),
        override_inventory=inspect_docx_overrides(output_docx),
    )
    val_errors = sum(1 for f in val_report.findings if f.severity == "error")
    val_warnings = sum(1 for f in val_report.findings if f.severity == "warning")

    # Step 6: Visual audit
    page_count = 0
    png_paths: list[Path] = []
    visual_finding_count = 0

    if include_visual_audit:
        logger.debug("Step 6: Visual audit")
        from .visual_audit import render_docx_to_png
        from .visual_audit_checks import run_visual_checks

        visual_dir = output_dir / "visual"
        render_report = render_docx_to_png(output_docx, visual_dir)
        if render_report.status == "pass":
            page_count = render_report.page_count
            png_paths = render_report.png_paths
            visual_findings = run_visual_checks(png_paths)
            visual_finding_count = len(visual_findings)

    status: Literal["success", "ai_failed", "build_failed", "validation_warnings"] = "success"
    if val_errors > 0:
        status = "validation_warnings"

    return GenerateReport(
        run_id=run_id, status=status, output_docx=output_docx,
        page_count=page_count, png_paths=png_paths,
        ai_model=resolved_model, ai_tokens=ai_tokens,
        validation_errors=val_errors, validation_warnings=val_warnings,
        visual_findings=visual_finding_count,
    )
