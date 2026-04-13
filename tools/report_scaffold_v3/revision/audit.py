from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph

from ..manifest_loader import load_baseline_manifest
from ..outline_audit import scan_docx_outline
from .common import compute_file_sha256, load_json_model, load_yaml_model, short_fingerprint, summarize_text
from .models import (
    IssueStatusRecommendation,
    OutlineRegressionSummary,
    RevisionAuditReport,
    RevisionExecutionReport,
    RevisionIssueLedger,
    RevisionPatchSpec,
    RevisionPostconditionResult,
)


def _paragraph_at(document: Document, index: Optional[int]):
    if index is None or index >= len(document.paragraphs):
        return None
    return document.paragraphs[index]


def _cell_at(
    document: Document,
    table_index: Optional[int],
    row_index: Optional[int],
    column_index: Optional[int],
):
    if table_index is None or row_index is None or column_index is None:
        return None
    if table_index >= len(document.tables):
        return None
    table = document.tables[table_index]
    if row_index >= len(table.rows):
        return None
    row = table.rows[row_index]
    if column_index >= len(row.cells):
        return None
    return row.cells[column_index]


def _find_operation_paragraph(document: Document, record, operation):
    matches = []
    target_text = None
    if operation.action == "replace_paragraph_text":
        target_text = operation.payload.get("text", "")
    if operation.action == "split_heading_and_body":
        target_text = operation.payload.get("heading_text", "")

    if target_text:
        expected = target_text.strip()
        matches = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip() == expected
        ]
    else:
        matches = [
            paragraph
            for paragraph in document.paragraphs
            if short_fingerprint(paragraph.text) == record.normalized_fingerprint
        ]

    if len(matches) == 1:
        return matches[0]

    return _paragraph_at(document, record.location.paragraph_index)


def _postcondition_results(
    candidate_path: Path,
    spec: RevisionPatchSpec,
) -> list[RevisionPostconditionResult]:
    document = Document(str(candidate_path))
    results: list[RevisionPostconditionResult] = []
    anchors = {anchor.anchor_id: anchor for anchor in spec.anchor_records}

    for operation in spec.operations:
        record = anchors[operation.anchor_id]
        messages: list[str] = []
        paragraph = _find_operation_paragraph(document, record, operation)
        cell = _cell_at(
            document,
            record.location.table_index,
            record.location.row_index,
            record.location.column_index,
        )
        for condition in operation.postconditions:
            if condition.kind == "paragraph_text_equals":
                actual = paragraph.text.strip() if paragraph is not None else ""
                if actual != (condition.value or ""):
                    messages.append(
                        f"{condition.kind} expected `{condition.value}`, got `{summarize_text(actual)}`."
                    )
            elif condition.kind == "paragraph_text_startswith":
                actual = paragraph.text.strip() if paragraph is not None else ""
                if not actual.startswith(condition.value or ""):
                    messages.append(
                        f"{condition.kind} expected prefix `{condition.value}`, got `{summarize_text(actual)}`."
                    )
            elif condition.kind == "paragraph_style_equals":
                actual_style = paragraph.style.name if paragraph is not None and paragraph.style else ""
                if actual_style != (condition.style or ""):
                    messages.append(
                        f"{condition.kind} expected `{condition.style}`, got `{actual_style}`."
                    )
            elif condition.kind == "inserted_paragraph_after_anchor":
                if paragraph is None:
                    messages.append(f"{condition.kind} could not resolve anchor paragraph.")
                    continue
                cursor = paragraph._p.getnext()
                found = False
                while cursor is not None:
                    if not cursor.tag.endswith("}p"):
                        cursor = cursor.getnext()
                        continue
                    probe = Paragraph(cursor, paragraph._parent)
                    if (probe.text or "").strip() == (condition.value or ""):
                        found = True
                        break
                    cursor = cursor.getnext()
                if not found:
                    messages.append(
                        f"{condition.kind} could not find inserted paragraph `{summarize_text(condition.value or '')}` after anchor."
                    )
            elif condition.kind == "table_cell_text_equals":
                actual = cell.text.strip() if cell is not None else ""
                if actual != (condition.value or ""):
                    messages.append(
                        f"{condition.kind} expected `{condition.value}`, got `{summarize_text(actual)}`."
                    )
        results.append(
            RevisionPostconditionResult(
                operation_id=operation.operation_id,
                issue_id=operation.issue_id,
                passed=not messages,
                messages=messages,
            )
        )
    return results


def _outline_regression(before_path: Path, after_path: Path) -> OutlineRegressionSummary:
    before = scan_docx_outline(before_path)
    after = scan_docx_outline(after_path)
    findings: list[str] = []

    if after.heading_count < before.heading_count:
        findings.append(
            f"Heading count regressed from {before.heading_count} to {after.heading_count}."
        )

    for level, count in before.heading_level_counts.items():
        if after.heading_level_counts.get(level, 0) < count:
            findings.append(
                f"Heading level {level} regressed from {count} to {after.heading_level_counts.get(level, 0)}."
            )

    missing_heading_texts = sorted({heading.text for heading in before.headings} - {heading.text for heading in after.headings})
    if missing_heading_texts:
        findings.append(
            "Missing pre-existing headings: " + ", ".join(missing_heading_texts[:5])
        )

    return OutlineRegressionSummary(
        has_regression=bool(findings),
        before_heading_count=before.heading_count,
        after_heading_count=after.heading_count,
        before_heading_level_counts=before.heading_level_counts,
        after_heading_level_counts=after.heading_level_counts,
        findings=findings,
    )


def _expected_paragraph_delta(spec: RevisionPatchSpec) -> int:
    delta = 0
    for operation in spec.operations:
        if operation.action == "insert_paragraph_after":
            delta += 1
        if operation.action == "split_heading_and_body":
            delta += 1
    return delta


def _non_target_drift(before_path: Path, after_path: Path, spec: RevisionPatchSpec) -> list[str]:
    before = Document(str(before_path))
    after = Document(str(after_path))
    findings: list[str] = []

    expected_paragraph_delta = _expected_paragraph_delta(spec)
    actual_paragraph_delta = len(after.paragraphs) - len(before.paragraphs)
    if actual_paragraph_delta != expected_paragraph_delta:
        findings.append(
            f"Paragraph count delta was {actual_paragraph_delta}, expected {expected_paragraph_delta}."
        )

    if len(after.tables) != len(before.tables):
        findings.append(
            f"Table count changed from {len(before.tables)} to {len(after.tables)}."
        )
    return findings


def build_revision_audit_report(
    candidate_path: Path,
    patch_spec_path: Path,
    execution_report_path: Path,
    issue_ledger_path: Path,
    baseline_manifest_path: Path,
) -> RevisionAuditReport:
    spec = load_yaml_model(patch_spec_path.expanduser().resolve(), RevisionPatchSpec)
    execution = load_json_model(execution_report_path.expanduser().resolve(), RevisionExecutionReport)
    ledger = load_yaml_model(issue_ledger_path.expanduser().resolve(), RevisionIssueLedger)
    baseline = load_baseline_manifest(baseline_manifest_path.expanduser().resolve())
    resolved_candidate = candidate_path.expanduser().resolve()

    if execution.backup_path is None:
        raise ValueError("Revision audit requires an execution report with a backup_path from a real apply run.")

    backup_path = execution.backup_path.expanduser().resolve()
    postconditions = _postcondition_results(resolved_candidate, spec)
    outline_regression = _outline_regression(backup_path, resolved_candidate)
    non_target_drift = _non_target_drift(backup_path, resolved_candidate, spec)

    protected_hash = compute_file_sha256(baseline.target_docx.expanduser().resolve())
    input_hash = execution.input_candidate_hash
    output_hash = compute_file_sha256(resolved_candidate)
    if output_hash == input_hash:
        non_target_drift.append("Candidate hash did not change after apply.")

    issue_ids = spec.target_issue_ids or [issue.issue_id for issue in ledger.issues]
    operation_pass = {result.operation_id: result.passed for result in postconditions}
    issue_pass = {
        issue_id: all(
            operation_pass.get(operation.operation_id, False)
            for operation in spec.operations
            if operation.issue_id == issue_id
        )
        for issue_id in issue_ids
    }

    overall_ok = (
        all(result.passed for result in postconditions)
        and not outline_regression.has_regression
        and not non_target_drift
    )

    fixed_issue_ids = sorted(issue_id for issue_id, passed in issue_pass.items() if passed and overall_ok)
    remaining_issue_ids = sorted(issue_id for issue_id in issue_ids if issue_id not in fixed_issue_ids)
    status_recommendations = [
        IssueStatusRecommendation(
            issue_id=issue_id,
            recommended_status="fixed" if issue_id in fixed_issue_ids else "verified",
            rationale=(
                "Revision audit passed for all bound operations."
                if issue_id in fixed_issue_ids
                else "Revision audit did not fully pass for this issue."
            ),
        )
        for issue_id in issue_ids
    ]

    return RevisionAuditReport(
        run_id=execution.run_id,
        target_issue_ids=issue_ids,
        input_candidate_hash=input_hash,
        output_candidate_hash=output_hash,
        protected_original_hash=protected_hash,
        postcondition_results=postconditions,
        outline_regression=outline_regression,
        non_target_drift_findings=non_target_drift,
        status="passed" if overall_ok else "failed",
        fixed_issue_ids=fixed_issue_ids,
        remaining_issue_ids=remaining_issue_ids,
        status_recommendations=status_recommendations,
    )


def render_revision_audit_markdown(report: RevisionAuditReport) -> str:
    lines = [
        "# Revision Audit Report",
        "",
        f"- Run id: `{report.run_id}`",
        f"- Target issue ids: {', '.join(report.target_issue_ids)}",
        f"- Input candidate hash: `{report.input_candidate_hash}`",
        f"- Output candidate hash: `{report.output_candidate_hash}`",
        f"- Protected original hash: `{report.protected_original_hash}`",
        f"- Status: `{report.status}`",
        "",
        "## Postcondition Results",
        "",
    ]

    for result in report.postcondition_results:
        prefix = "passed" if result.passed else "failed"
        lines.append(f"- `{result.operation_id}` / `{prefix}`")
        for message in result.messages:
            lines.append(f"  - {message}")

    lines.extend(["", "## Outline Regression", ""])
    if not report.outline_regression.findings:
        lines.append("- No outline regression detected.")
    else:
        for finding in report.outline_regression.findings:
            lines.append(f"- {finding}")

    lines.extend(["", "## Non-target Drift", ""])
    if not report.non_target_drift_findings:
        lines.append("- No non-target drift detected.")
    else:
        for finding in report.non_target_drift_findings:
            lines.append(f"- {finding}")

    lines.extend(["", "## Status Recommendations", ""])
    for recommendation in report.status_recommendations:
        lines.append(
            f"- `{recommendation.issue_id}` -> `{recommendation.recommended_status}`: {recommendation.rationale}"
        )
    return "\n".join(lines)
