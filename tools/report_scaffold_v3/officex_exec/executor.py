from __future__ import annotations

from dataclasses import dataclass
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from ..docx_inspector import normalize_text
from ..runtime_common import load_runtime_structured_model
from .common import compute_file_sha256, make_exec_id, short_fingerprint, summarize_text
from .models import (
    OfficeXExecutionOperation,
    OfficeXExecutionPatchSpec,
    OfficeXExecutionReport,
    OfficeXLiveAnchorRecord,
    OfficeXOperationExecutionResult,
)


@dataclass
class ResolvedAnchor:
    record: OfficeXLiveAnchorRecord
    paragraph: Paragraph | None = None
    cell: _Cell | None = None


@dataclass
class OperationRuntime:
    operation: OfficeXExecutionOperation
    anchor: ResolvedAnchor
    inserted_paragraph: Paragraph | None = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    if text:
        new_para.add_run(text)
    return new_para


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    p = paragraph._p
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    paragraph.add_run(text)


def _fingerprint(text: str) -> str:
    return short_fingerprint(text)


def _resolve_paragraph_anchor(document: Document, record: OfficeXLiveAnchorRecord) -> ResolvedAnchor:
    index = record.location.paragraph_index
    if index is None or index >= len(document.paragraphs):
        raise ValueError(f"Anchor `{record.anchor_id}` points to invalid paragraph index `{index}`.")
    paragraph = document.paragraphs[index]
    fingerprint = normalize_text(paragraph.text)
    if not fingerprint:
        raise ValueError(f"Anchor `{record.anchor_id}` resolved to an empty paragraph.")
    if record.normalized_fingerprint != _fingerprint(fingerprint):
        raise ValueError(
            f"Anchor `{record.anchor_id}` fingerprint mismatch at {record.location_hint}; expected `{record.normalized_fingerprint}`."
        )
    return ResolvedAnchor(record=record, paragraph=paragraph)


def _resolve_table_anchor(document: Document, record: OfficeXLiveAnchorRecord) -> ResolvedAnchor:
    table_index = record.location.table_index
    row_index = record.location.row_index
    column_index = record.location.column_index
    if table_index is None or row_index is None or column_index is None:
        raise ValueError(f"Anchor `{record.anchor_id}` is missing table coordinates.")
    if table_index >= len(document.tables):
        raise ValueError(f"Anchor `{record.anchor_id}` points to invalid table index `{table_index}`.")
    table = document.tables[table_index]
    if row_index >= len(table.rows) or column_index >= len(table.rows[row_index].cells):
        raise ValueError(f"Anchor `{record.anchor_id}` points to invalid cell coordinates.")
    cell = table.rows[row_index].cells[column_index]
    if record.normalized_fingerprint != _fingerprint(cell.text):
        raise ValueError(
            f"Anchor `{record.anchor_id}` fingerprint mismatch at {record.location_hint}; expected `{record.normalized_fingerprint}`."
        )
    return ResolvedAnchor(record=record, cell=cell)


def _resolve_anchor_map(document: Document, spec: OfficeXExecutionPatchSpec) -> dict[str, ResolvedAnchor]:
    mapping: dict[str, ResolvedAnchor] = {}
    for record in spec.anchor_records:
        if record.uniqueness_status != "unique":
            raise ValueError(
                f"Anchor `{record.anchor_id}` is not executable because status is `{record.uniqueness_status}`."
            )
        if record.block_kind in {"paragraph", "heading"}:
            mapping[record.anchor_id] = _resolve_paragraph_anchor(document, record)
        elif record.block_kind == "table_cell":
            mapping[record.anchor_id] = _resolve_table_anchor(document, record)
        else:
            raise ValueError(f"Unsupported OfficeX anchor block kind `{record.block_kind}`.")
    return mapping


def _apply_operation(runtime: OperationRuntime) -> None:
    operation = runtime.operation
    resolved = runtime.anchor

    if operation.action == "insert_paragraph_after":
        if resolved.paragraph is None:
            raise ValueError(f"Operation `{operation.operation_id}` requires a paragraph anchor.")
        runtime.inserted_paragraph = insert_paragraph_after(
            resolved.paragraph,
            operation.payload.get("text", ""),
            operation.payload.get("style"),
        )
        return

    if operation.action == "replace_paragraph_text":
        if resolved.paragraph is None:
            raise ValueError(f"Operation `{operation.operation_id}` requires a paragraph anchor.")
        replace_paragraph_text(resolved.paragraph, operation.payload.get("text", ""))
        return

    if operation.action == "update_paragraph_style":
        if resolved.paragraph is None:
            raise ValueError(f"Operation `{operation.operation_id}` requires a paragraph anchor.")
        resolved.paragraph.style = operation.payload.get("style", "")
        return

    if operation.action == "update_table_cell":
        if resolved.cell is None:
            raise ValueError(f"Operation `{operation.operation_id}` requires a table-cell anchor.")
        updates = operation.payload.get("table_cell_updates", [])
        if len(updates) != 1:
            raise ValueError(
                f"Operation `{operation.operation_id}` expects exactly one table_cell_updates entry."
            )
        resolved.cell.text = updates[0].get("text", "")
        return

    raise ValueError(f"Unsupported OfficeX execution action `{operation.action}`.")


def _check_postconditions(runtime: OperationRuntime) -> list[str]:
    messages: list[str] = []
    operation = runtime.operation
    resolved = runtime.anchor

    for condition in operation.postconditions:
        if condition.kind == "paragraph_text_equals":
            actual = normalize_text(resolved.paragraph.text if resolved.paragraph is not None else "")
            expected = normalize_text(condition.value or "")
            if actual != expected:
                messages.append(f"{condition.kind} failed: expected `{expected}`, got `{actual}`.")
        elif condition.kind == "paragraph_style_equals":
            actual_style = resolved.paragraph.style.name if resolved.paragraph is not None and resolved.paragraph.style else ""
            if actual_style != (condition.style or ""):
                messages.append(f"{condition.kind} failed: expected `{condition.style}`, got `{actual_style}`.")
        elif condition.kind == "inserted_paragraph_after_anchor":
            if runtime.inserted_paragraph is None or resolved.paragraph is None:
                messages.append(f"{condition.kind} failed: no inserted paragraph recorded.")
                continue
            inserted_text = normalize_text(runtime.inserted_paragraph.text)
            expected = normalize_text(condition.value or "")
            if inserted_text != expected:
                messages.append(f"{condition.kind} failed: expected inserted text `{expected}`, got `{inserted_text}`.")
            if condition.style:
                actual_style = runtime.inserted_paragraph.style.name if runtime.inserted_paragraph.style else ""
                if actual_style != condition.style:
                    messages.append(
                        f"{condition.kind} failed: expected inserted style `{condition.style}`, got `{actual_style}`."
                    )
            cursor = resolved.paragraph._p.getnext()
            found = False
            while cursor is not None:
                if cursor is runtime.inserted_paragraph._p:
                    found = True
                    break
                cursor = cursor.getnext()
            if not found:
                messages.append(f"{condition.kind} failed: inserted paragraph is not after the anchor.")
        elif condition.kind == "table_cell_text_equals":
            actual = normalize_text(resolved.cell.text if resolved.cell is not None else "")
            expected = normalize_text(condition.value or "")
            if actual != expected:
                messages.append(f"{condition.kind} failed: expected `{expected}`, got `{actual}`.")
        else:
            messages.append(f"Unsupported postcondition kind `{condition.kind}`.")
    return messages


def _result_target_summary(anchor: OfficeXLiveAnchorRecord) -> str:
    return f"{anchor.location_hint} | {anchor.block_kind} | {anchor.text_excerpt or '[none]'}"


def execute_officex_patch_spec(
    candidate_path: Path,
    patch_spec_path: Path,
    *,
    dry_run: bool,
    backup_dir: Path | None = None,
) -> OfficeXExecutionReport:
    resolved_candidate = candidate_path.expanduser().resolve()
    resolved_patch_spec = patch_spec_path.expanduser().resolve()
    spec = load_runtime_structured_model(resolved_patch_spec, OfficeXExecutionPatchSpec)
    if resolved_candidate != spec.candidate_path:
        raise ValueError(
            f"Refusing to apply OfficeX spec for `{spec.candidate_path}` to different candidate `{resolved_candidate}`."
        )

    input_hash = compute_file_sha256(resolved_candidate)
    if input_hash != spec.candidate_hash:
        raise ValueError(
            f"Candidate hash mismatch for `{resolved_candidate}`: expected `{spec.candidate_hash}`, got `{input_hash}`."
        )

    report = OfficeXExecutionReport(
        run_id=make_exec_id("officex-apply-dry-run" if dry_run else "officex-apply"),
        spec_id=spec.spec_id,
        input_candidate=resolved_candidate,
        input_candidate_hash=input_hash,
        spec_path=resolved_patch_spec,
        dry_run=dry_run,
        status="validated" if dry_run else "applied",
    )

    document = Document(str(resolved_candidate))
    anchor_map = _resolve_anchor_map(document, spec)
    runtimes: list[OperationRuntime] = []
    try:
        for operation in spec.operations:
            if operation.expected_candidate_hash != input_hash:
                raise ValueError(
                    f"Operation `{operation.operation_id}` expected hash `{operation.expected_candidate_hash}`, got `{input_hash}`."
                )
            if operation.anchor_id not in anchor_map:
                raise ValueError(f"Operation `{operation.operation_id}` references unknown anchor `{operation.anchor_id}`.")
            runtime = OperationRuntime(operation=operation, anchor=anchor_map[operation.anchor_id])
            _apply_operation(runtime)
            failures = _check_postconditions(runtime)
            if failures:
                raise ValueError(
                    f"Operation `{operation.operation_id}` failed postconditions: {'; '.join(failures)}"
                )
            runtimes.append(runtime)
            report.applied_operations.append(
                OfficeXOperationExecutionResult(
                    operation_id=operation.operation_id,
                    issue_id=operation.issue_id,
                    anchor_id=operation.anchor_id,
                    action=operation.action,
                    status="validated" if dry_run else "applied",
                    target_summary=_result_target_summary(runtime.anchor.record),
                    postconditions_checked=len(operation.postconditions),
                )
            )
        if dry_run:
            report.output_candidate_hash = None
            return report

        if backup_dir is None:
            backup_dir = resolved_candidate.parent
        backup_dir.mkdir(parents=True, exist_ok=True)
        backup_path = backup_dir / f"{resolved_candidate.stem}.{spec.spec_id}.backup{resolved_candidate.suffix}"
        shutil.copy2(resolved_candidate, backup_path)
        report.backup_path = backup_path
        document.save(str(resolved_candidate))
        report.output_candidate_hash = compute_file_sha256(resolved_candidate)
        return report
    except Exception as exc:
        report.status = "failed"
        report.failure_reason = str(exc)
        if runtimes:
            last = runtimes[-1]
            report.rejected_operations.append(
                OfficeXOperationExecutionResult(
                    operation_id=last.operation.operation_id,
                    issue_id=last.operation.issue_id,
                    anchor_id=last.operation.anchor_id,
                    action=last.operation.action,
                    status="rejected",
                    target_summary=_result_target_summary(last.anchor.record),
                    postconditions_checked=len(last.operation.postconditions),
                    reason=str(exc),
                )
            )
        return report


def render_officex_execution_markdown(report: OfficeXExecutionReport) -> str:
    lines = [
        "# OfficeX Execution Report",
        "",
        f"- Run id: `{report.run_id}`",
        f"- Spec id: `{report.spec_id}`",
        f"- Input candidate: `{report.input_candidate}`",
        f"- Input candidate hash: `{report.input_candidate_hash}`",
        f"- Spec path: `{report.spec_path}`",
        f"- Dry run: `{report.dry_run}`",
        f"- Status: `{report.status}`",
        f"- Backup path: `{report.backup_path}`" if report.backup_path else "- Backup path: `[none]`",
        f"- Output candidate hash: `{report.output_candidate_hash}`" if report.output_candidate_hash else "- Output candidate hash: `[none]`",
        f"- Failure reason: {report.failure_reason}" if report.failure_reason else "- Failure reason: [none]",
        "",
        "## Applied Operations",
        "",
    ]

    if not report.applied_operations:
        lines.append("- No operations applied.")
    else:
        for result in report.applied_operations:
            lines.append(
                f"- `{result.operation_id}` / `{result.issue_id}` / `{result.action}` / `{result.status}` / `{result.target_summary}`"
            )

    lines.extend(["", "## Rejected Operations", ""])
    if not report.rejected_operations:
        lines.append("- No rejected operations.")
    else:
        for result in report.rejected_operations:
            lines.append(
                f"- `{result.operation_id}` / `{result.issue_id}` / `{result.action}` / `{result.status}` / {result.reason or '[none]'}"
            )
    return "\n".join(lines)
