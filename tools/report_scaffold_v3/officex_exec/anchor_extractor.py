from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

from ..docx_inspector import normalize_text
from ..models import OfficeXAnchorRule, OfficeXParagraphAnchorRule, OfficeXTableCellAnchorRule
from .common import compute_file_sha256, make_exec_id, render_context_lines, short_fingerprint, summarize_text, utc_now_iso
from .models import OfficeXAnchorLocation, OfficeXLiveAnchorRecord, OfficeXLiveAnchorSnapshot


HEADING_STYLE_ID_PATTERN = re.compile(r"^heading\s*(\d+)$", re.IGNORECASE)


@dataclass(frozen=True)
class ParagraphScanEntry:
    index: int
    paragraph: object
    normalized_text: str
    raw_style_id: str | None
    is_heading: bool


def _raw_paragraph_style_id(paragraph: object) -> str | None:
    paragraph_element = getattr(paragraph, "_p", None)
    paragraph_properties = getattr(paragraph_element, "pPr", None)
    raw_style_id = getattr(paragraph_properties, "style", None) if paragraph_properties is not None else None
    if raw_style_id is None:
        return None
    value = str(raw_style_id).strip()
    return value or None


def _is_heading_style_id(style_id: str | None) -> bool:
    if not style_id:
        return False
    return HEADING_STYLE_ID_PATTERN.match(style_id) is not None


def _paragraph_style_name(paragraph: object, raw_style_id: str | None) -> str | None:
    style = getattr(paragraph, "style", None)
    if style is not None:
        style_name = getattr(style, "name", None)
        if style_name:
            return str(style_name)
    return raw_style_id


def _build_paragraph_scan_entries(document: Document) -> list[ParagraphScanEntry]:
    entries: list[ParagraphScanEntry] = []
    for index, paragraph in enumerate(document.paragraphs):
        raw_style_id = _raw_paragraph_style_id(paragraph)
        entries.append(
            ParagraphScanEntry(
                index=index,
                paragraph=paragraph,
                normalized_text=normalize_text(paragraph.text),
                raw_style_id=raw_style_id,
                is_heading=_is_heading_style_id(raw_style_id),
            )
        )
    return entries


def _non_empty_context(entries: list[ParagraphScanEntry], index: int, *, direction: int, limit: int = 2) -> list[str]:
    values: list[str] = []
    cursor = index + direction
    while 0 <= cursor < len(entries) and len(values) < limit:
        text = entries[cursor].normalized_text
        if text:
            values.append(summarize_text(text, limit=100))
        cursor += direction
    if direction < 0:
        values.reverse()
    return values


def _anchor_id(issue_id: str, anchor_role: str, location_hint: str, text: str) -> str:
    return f"{issue_id}:{anchor_role}:{location_hint}:{short_fingerprint(text)}"


def _build_paragraph_anchor(
    *,
    issue_id: str,
    anchor_role: str,
    candidate_path: Path,
    candidate_hash: str,
    block_kind: str,
    location_hint: str,
    paragraph_index: int | None,
    text: str,
    style_name: str | None,
    leading_context: list[str],
    trailing_context: list[str],
    uniqueness_status: str,
    match_count: int,
) -> OfficeXLiveAnchorRecord:
    anchor_text = text if text else f"{issue_id}:{anchor_role}:{location_hint}"
    return OfficeXLiveAnchorRecord(
        candidate_path=candidate_path,
        candidate_hash=candidate_hash,
        issue_id=issue_id,
        anchor_id=_anchor_id(issue_id, anchor_role, location_hint, anchor_text),
        anchor_role=anchor_role,
        block_kind=block_kind,
        location_hint=location_hint,
        location=OfficeXAnchorLocation(block_kind=block_kind, paragraph_index=paragraph_index),
        normalized_fingerprint=short_fingerprint(anchor_text),
        uniqueness_status=uniqueness_status,  # type: ignore[arg-type]
        text_excerpt=summarize_text(text),
        style_name=style_name,
        leading_context=leading_context,
        trailing_context=trailing_context,
        match_count=match_count,
    )


def _extract_paragraph_anchor(
    paragraph_entries: list[ParagraphScanEntry],
    *,
    candidate_path: Path,
    candidate_hash: str,
    rule: OfficeXParagraphAnchorRule,
) -> OfficeXLiveAnchorRecord:
    needle = normalize_text(rule.needle)
    matches: list[ParagraphScanEntry] = []

    for entry in paragraph_entries:
        text = entry.normalized_text
        if not text:
            continue
        if rule.block_kind == "heading" and not entry.is_heading:
            continue
        if rule.match_mode == "exact" and text == needle:
            matches.append(entry)
        if rule.match_mode == "prefix" and text.startswith(needle):
            matches.append(entry)

    if len(matches) == 1:
        entry = matches[0]
        return _build_paragraph_anchor(
            issue_id=rule.issue_id,
            anchor_role=rule.anchor_role,
            candidate_path=candidate_path,
            candidate_hash=candidate_hash,
            block_kind=rule.block_kind,
            location_hint=f"P{entry.index}",
            paragraph_index=entry.index,
            text=entry.normalized_text,
            style_name=_paragraph_style_name(entry.paragraph, entry.raw_style_id),
            leading_context=_non_empty_context(paragraph_entries, entry.index, direction=-1),
            trailing_context=_non_empty_context(paragraph_entries, entry.index, direction=1),
            uniqueness_status="unique",
            match_count=1,
        )

    if not matches:
        return _build_paragraph_anchor(
            issue_id=rule.issue_id,
            anchor_role=rule.anchor_role,
            candidate_path=candidate_path,
            candidate_hash=candidate_hash,
            block_kind=rule.block_kind,
            location_hint="missing",
            paragraph_index=None,
            text="",
            style_name=None,
            leading_context=[],
            trailing_context=[],
            uniqueness_status="missing",
            match_count=0,
        )

    indices = [entry.index for entry in matches]
    first = matches[0]
    return _build_paragraph_anchor(
        issue_id=rule.issue_id,
        anchor_role=rule.anchor_role,
        candidate_path=candidate_path,
        candidate_hash=candidate_hash,
        block_kind=rule.block_kind,
        location_hint=f"multiple:{','.join(f'P{value}' for value in indices)}",
        paragraph_index=first.index,
        text=first.normalized_text,
        style_name=_paragraph_style_name(first.paragraph, first.raw_style_id),
        leading_context=_non_empty_context(paragraph_entries, first.index, direction=-1),
        trailing_context=_non_empty_context(paragraph_entries, first.index, direction=1),
        uniqueness_status="non_unique",
        match_count=len(matches),
    )


def _extract_table_cell_anchor(
    document: Document,
    *,
    candidate_path: Path,
    candidate_hash: str,
    rule: OfficeXTableCellAnchorRule,
) -> OfficeXLiveAnchorRecord:
    header_values = tuple(normalize_text(value) for value in rule.header_values)
    matches: list[tuple[int, int, int, object, list[str], list[str]]] = []

    for table_index, table in enumerate(document.tables):
        if not table.rows:
            continue
        table_header = tuple(normalize_text(cell.text) for cell in table.rows[0].cells)
        if table_header != header_values:
            continue
        for row_index, row in enumerate(table.rows[1:], start=1):
            cells = [normalize_text(cell.text) for cell in row.cells]
            if len(cells) <= max(rule.row_key_column, rule.target_column):
                continue
            if cells[rule.row_key_column] != normalize_text(rule.row_key):
                continue
            target_cell = row.cells[rule.target_column]
            previous_row: list[str] = []
            next_row: list[str] = []
            if row_index - 1 >= 1:
                previous_row = [" | ".join(normalize_text(cell.text) for cell in table.rows[row_index - 1].cells)]
            if row_index + 1 < len(table.rows):
                next_row = [" | ".join(normalize_text(cell.text) for cell in table.rows[row_index + 1].cells)]
            matches.append((table_index, row_index, rule.target_column, target_cell, previous_row, next_row))

    if len(matches) == 1:
        table_index, row_index, column_index, cell, previous_row, next_row = matches[0]
        cell_text = normalize_text(cell.text)
        location_hint = f"T{table_index}:R{row_index}:C{column_index}"
        return OfficeXLiveAnchorRecord(
            candidate_path=candidate_path,
            candidate_hash=candidate_hash,
            issue_id=rule.issue_id,
            anchor_id=_anchor_id(rule.issue_id, rule.anchor_role, location_hint, cell_text),
            anchor_role=rule.anchor_role,
            block_kind="table_cell",
            location_hint=location_hint,
            location=OfficeXAnchorLocation(
                block_kind="table_cell",
                table_index=table_index,
                row_index=row_index,
                column_index=column_index,
            ),
            normalized_fingerprint=short_fingerprint(cell_text),
            uniqueness_status="unique",
            text_excerpt=summarize_text(cell_text),
            style_name=None,
            leading_context=previous_row,
            trailing_context=next_row,
            match_count=1,
        )

    if not matches:
        return OfficeXLiveAnchorRecord(
            candidate_path=candidate_path,
            candidate_hash=candidate_hash,
            issue_id=rule.issue_id,
            anchor_id=_anchor_id(rule.issue_id, rule.anchor_role, "missing", rule.anchor_role),
            anchor_role=rule.anchor_role,
            block_kind="table_cell",
            location_hint="missing",
            location=OfficeXAnchorLocation(block_kind="table_cell"),
            normalized_fingerprint=short_fingerprint(rule.anchor_role),
            uniqueness_status="missing",
            text_excerpt="",
            style_name=None,
            leading_context=[],
            trailing_context=[],
            match_count=0,
        )

    table_index, row_index, column_index, cell, previous_row, next_row = matches[0]
    location_hint = "multiple:" + ",".join(
        f"T{table_idx}:R{row_idx}:C{col_idx}" for table_idx, row_idx, col_idx, *_ in matches
    )
    cell_text = normalize_text(cell.text)
    return OfficeXLiveAnchorRecord(
        candidate_path=candidate_path,
        candidate_hash=candidate_hash,
        issue_id=rule.issue_id,
        anchor_id=_anchor_id(rule.issue_id, rule.anchor_role, location_hint, cell_text),
        anchor_role=rule.anchor_role,
        block_kind="table_cell",
        location_hint=location_hint,
        location=OfficeXAnchorLocation(
            block_kind="table_cell",
            table_index=table_index,
            row_index=row_index,
            column_index=column_index,
        ),
        normalized_fingerprint=short_fingerprint(cell_text),
        uniqueness_status="non_unique",
        text_excerpt=summarize_text(cell_text),
        style_name=None,
        leading_context=previous_row,
        trailing_context=next_row,
        match_count=len(matches),
    )


def build_officex_live_anchor_snapshot(
    candidate_path: Path,
    review_ledger_path: Path,
    issue_ids: list[str],
    anchor_rules: list[OfficeXAnchorRule],
) -> tuple[OfficeXLiveAnchorSnapshot, list[str]]:
    resolved_candidate = candidate_path.expanduser().resolve()
    candidate_hash = compute_file_sha256(resolved_candidate)
    document = Document(str(resolved_candidate))
    paragraph_entries = _build_paragraph_scan_entries(document)
    anchors: list[OfficeXLiveAnchorRecord] = []
    findings: list[str] = []

    for rule in anchor_rules:
        if isinstance(rule, OfficeXParagraphAnchorRule):
            anchor = _extract_paragraph_anchor(
                paragraph_entries,
                candidate_path=resolved_candidate,
                candidate_hash=candidate_hash,
                rule=rule,
            )
        else:
            anchor = _extract_table_cell_anchor(
                document,
                candidate_path=resolved_candidate,
                candidate_hash=candidate_hash,
                rule=rule,
            )
        anchors.append(anchor)
        if anchor.uniqueness_status != "unique":
            findings.append(
                f"{anchor.issue_id}:{anchor.anchor_role} resolved as {anchor.uniqueness_status} "
                f"(match_count={anchor.match_count}, location={anchor.location_hint})."
            )

    snapshot = OfficeXLiveAnchorSnapshot(
        snapshot_id=make_exec_id("officex-live-anchor-snapshot"),
        generated_at_utc=utc_now_iso(),
        candidate_path=resolved_candidate,
        candidate_hash=candidate_hash,
        source_review_ledger_path=review_ledger_path.expanduser().resolve(),
        target_issue_ids=issue_ids,
        anchors=anchors,
    )
    return snapshot, findings


def render_officex_live_anchor_snapshot_markdown(
    snapshot: OfficeXLiveAnchorSnapshot,
    findings: list[str],
) -> str:
    lines = [
        "# OfficeX Live Anchor Snapshot",
        "",
        f"- Snapshot id: `{snapshot.snapshot_id}`",
        f"- Generated at (UTC): `{snapshot.generated_at_utc}`",
        f"- Candidate path: `{snapshot.candidate_path}`",
        f"- Candidate hash: `{snapshot.candidate_hash}`",
        f"- Review ledger: `{snapshot.source_review_ledger_path}`",
        f"- Target issue ids: {', '.join(snapshot.target_issue_ids)}",
        "",
        "## Anchors",
        "",
    ]

    for anchor in snapshot.anchors:
        lines.extend(
            [
                f"- `{anchor.anchor_id}`",
                f"  issue `{anchor.issue_id}` / role `{anchor.anchor_role}` / kind `{anchor.block_kind}` / status `{anchor.uniqueness_status}`",
                f"  location `{anchor.location_hint}` / style `{anchor.style_name or '[none]'}` / fingerprint `{anchor.normalized_fingerprint}`",
                f"  excerpt: {anchor.text_excerpt or '[none]'}",
                f"  leading: {render_context_lines(anchor.leading_context)}",
                f"  trailing: {render_context_lines(anchor.trailing_context)}",
            ]
        )

    lines.extend(["", "## Findings", ""])
    if not findings:
        lines.append("- All requested anchors resolved uniquely.")
    else:
        for finding in findings:
            lines.append(f"- {finding}")
    return "\n".join(lines)
